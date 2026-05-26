from __future__ import annotations

from pathlib import Path

from docx import Document


PATHS = [
    Path(r"C:\Users\Twist\Desktop\论文\主稿\毕业论文初稿v2.8_模板格式终修版_2026-04-06.docx"),
    Path(r"C:\Users\Twist\Desktop\论文\主稿\毕业论文初稿v2.8_降风险与引用终修版_v2_2026-04-06.docx"),
    Path(r"C:\Users\Twist\Desktop\论文\主稿\毕业论文初稿v2.8_降风险与引用终修版_v3_2026-04-06.docx"),
]


def main() -> None:
    for path in PATHS:
        print("FILE", path)
        print("EXISTS", path.exists())
        if not path.exists():
            print("---")
            continue

        doc = Document(str(path))
        nonempty = [(i, p.text.strip()) for i, p in enumerate(doc.paragraphs) if p.text.strip()]
        print("PARAS", len(doc.paragraphs), "NONEMPTY", len(nonempty), "TABLES", len(doc.tables))
        print("HAS_QMARK", any("????" in text for _, text in nonempty))
        for i, text in nonempty[:20]:
            print(i, text[:120])
        print("---")


if __name__ == "__main__":
    main()
