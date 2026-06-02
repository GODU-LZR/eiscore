# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (c) 2026 林志荣

from docx import Document
from pathlib import Path


desktop = Path(r"C:\Users\Twist\Desktop\论文")
out_dir = desktop / "清单说明"
out_dir.mkdir(parents=True, exist_ok=True)

md_path = out_dir / "v2.8_数据库设计重写说明与模拟AIGC结论_2026-04-12.md"
docx_path = out_dir / "v2.8_数据库设计重写说明与模拟AIGC结论_2026-04-12.docx"

content = """# v2.8 数据库设计重写说明与模拟 AIGC 结论

## 一、本次第五章实际修改内容

本次修改不是局部润色，而是对第五章“数据库设计”进行了整章重写，重点解决了原稿中存在的以下问题：

1. 数据库设计目标、总体结构、核心实体关系、关键物理表结构之间的层次混乱。
2. 部分段落把“表说明”“章节说明”和“字段表内容”混在一起，阅读时容易断裂。
3. `5.7`、`5.8`、`5.9` 一带存在重复引入、跳段和孤立标题问题。
4. 第五章虽然提到实体关系，但缺少能够支撑正文叙述的 ER 图落点。
5. 关键表结构前后的正文引入和解释不足，容易让老师觉得是在堆表格。

针对这些问题，当前版本做了以下重写：

- 重新组织为 `5.1` 到 `5.10` 的清晰结构。
- 保留并补齐 `图5-1 EISCore 核心数据实体关系图`。
- 把组织权限、物料仓储、应用中心与流程、数据安全与权限控制、语义增强等部分重新分层。
- 为 `public.users`、`public.roles`、`public.user_roles`、`public.raw_materials`、`scm.warehouses`、`scm.inventory_batches`、`scm.inventory_transactions`、`app_center.apps`、`workflow.instances` 九张关键物理表逐一补了正文引入句和解释段。
- 删除了原稿中重复、错位、孤立的小节内容，使第五章更接近“先说明设计，再展示结构，再落到关键表”的论文写法。

## 二、按照老师批注要求的落实情况

从老师前面的批注方向看，第五章最相关的要求主要是：

1. 不能只堆表格，图和表都要有正文引用与说明。
2. 标题和正文不能混写，尤其不能出现“标题像正文、正文像标题”的情况。
3. 章节结构要清楚，图表顺序、说明顺序要能让读者顺着读下去。

这一版第五章的落实情况如下：

- `图5-1` 已经在正文中引入，并在后文围绕三条主线展开解释。
- 每张关键物理表前都先交代“为什么需要这张表”以及“该表在系统中的职责”，不是直接摆字段。
- 小节标题统一成简洁的章节式写法，例如“用户表”“角色表”“库存批次表”，表名放回正文，不再把标题和正文混在一起。
- `5.10 本章小结` 重新收束到“数据库不仅承担存储，也承担权限、流程、语义治理”的整体结论，与课题主线一致。

## 三、本地启发式检测结果

扫描对象：

- `毕业论文初稿v2.8_数据库设计重写版_2026-04-12.docx`

本地启发式扫描结果：

- 段落数：`65`
- 平均 AI 风险分：`0.2171`
- 高风险段落：`5`
- 中风险段落：`8`
- 低风险段落：`52`

需要说明的是，这个分数反映的是**整份主稿**，不只是第五章本身。当前被拉高的主要位置仍然集中在：

- 英文 `Introduction`
- 第二章本章小结及概括性技术段
- 第四章/第六章中的高概括设计总述段

也就是说，这次第五章重写并不是当前高风险的主要来源。

## 四、本地模拟 AIGC 检测结论

这里的“模拟检测”不是学校或第三方官方系统结果，而是依据项目目录中已有的两类材料进行规则式判断：

- `AIGC模拟检测结论_2026-03-15.md`
- `AIGC检测报告分析与合规修改建议_2026-03-15.md`

按这些文档的口径，当前第五章数据库设计部分的表现总体上比旧稿更稳，原因主要有三点：

1. 这一版减少了“该表用于……该表用于……”式重复句，降低了表格式模板感。
2. 文本围绕 `Schema` 划分、实体关系、批次链路、流程状态写回、RLS 和语义增强展开，项目锚点更明显。
3. ER 图与物理表结构之间形成了“先总后分”的叙述关系，不再是孤立的表格堆叠。

按模拟检测口径，这一版第五章的判断可归为：

- **章节层面：中低风险**
- **风险主要来源：不是第五章，而是整篇中仍然偏概括的前置章节和英摘部分**

## 五、当前版本还值得继续优化的地方

如果后面还要继续收这一章，可以优先做下面三件事：

1. 在 `5.2` 和 `5.3` 各再补一两句“为什么采用多 Schema / 为什么三条主线这样组织”的工程取舍说明。
2. 在 `5.7` 数据安全与权限控制设计中，再增加一两句与第六章权限链路的呼应。
3. 在 `5.8` 语义增强相关设计中，再补一句它与应用中心、流程和动态数据应用之间的边界关系。

整体上，第五章当前已经比旧稿明显更像“数据库设计章节”，而不是“表结构堆叠章节”。
"""

md_path.write_text(content, encoding="utf-8")

doc = Document()
for block in content.split("\n\n"):
    lines = block.splitlines()
    if lines and lines[0].startswith("# "):
        p = doc.add_paragraph()
        p.style = "Heading 1"
        p.add_run(lines[0][2:])
        for line in lines[1:]:
            if line.strip():
                doc.add_paragraph(line)
    elif lines and lines[0].startswith("## "):
        p = doc.add_paragraph()
        p.style = "Heading 2"
        p.add_run(lines[0][3:])
        for line in lines[1:]:
            if line.strip():
                doc.add_paragraph(line)
    else:
        for line in lines:
            doc.add_paragraph(line)

doc.save(docx_path)
print(f"saved: {md_path}")
print(f"saved: {docx_path}")
