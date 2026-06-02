# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (c) 2026 林志荣

from docx import Document
from pathlib import Path


INPUT = Path(r"\\wsl.localhost\Ubuntu\home\lzr\eiscore\docs\tmp_thesis_v2_for_v21.docx")
OUTPUT = Path(r"\\wsl.localhost\Ubuntu\home\lzr\eiscore\docs\毕业论文初稿v2.1_定向优化版_2026-03-15.docx")
OUTPUT_EN = Path(r"\\wsl.localhost\Ubuntu\home\lzr\eiscore\docs\thesis_draft_v2_1_targeted_2026-03-15.docx")


REWRITE = {
    197: "在南派食品的业务调研里，最先暴露出来的并不是“系统功能够不够多”，而是同一批业务数据在纸单、微信群消息和分散台账之间反复转手。销售订单、领料单、检验附件和库存记录只要断开一处，后面再追某批原料进了哪个仓位、对应哪次补料、最终流向哪张出库单，速度就会明显慢下来。问题就出在这里。湛江不少中小制造企业都面临类似约束：预算不宽，IT 运维力量薄弱，业务规则又会随着订单、原料和现场安排持续变化。对这类企业来说，真正缺的往往不是再上一套更重的系统，而是一套能先把物料、人员、流程和数据关系接稳的底座。",
    198: "本文据此设计并实现了 EISCore。系统没有把主要精力放在继续堆叠中间服务，而是先把核心数据、部分业务约束和访问边界收束到 PostgreSQL 与 PostgREST 这条链路，再由微前端基座承载人事、物料、应用中心和移动端。数据库、接口、运行时和前端进程统一交给 Docker Compose 与 PM2 组织。这样取舍并不花哨，但更适合当前场景：先把人事治理、库存台账、应用配置、流程联动和语义增强几条最容易失控的链路收稳，再谈后续扩展。",
    202: "For Nanpai Food, digitalization becomes difficult at the point where one business record is split into several manual traces. Orders, lot records, inspection files and stock movements still travel across paper forms, chat messages and separate ledgers. Once one step is delayed or copied twice, traceability becomes slower and coordination becomes unreliable. This is a common constraint in resource-constrained manufacturers in Zhanjiang, where business changes faster than the available IT maintenance capacity.",
    203: "This thesis therefore develops EISCore as a lightweight system core rather than a full-scale enterprise suite. The project does not start from a heavy service layer. Instead, PostgreSQL carries the core data model together with part of the business rules and access boundaries, while PostgREST reduces repetitive API work. On the client side, Vue 3 and qiankun keep one entry while separating the modules that have actually been stabilized, including human resource management, material management, the application center and mobile services. Workflow support, FlashBuilder, controlled agent tools, lightweight semantic enhancement and the cold-storage PDA mode are then organized around the same runtime chain.",
    204: "The thesis follows the actual engineering path of the project. Chapter 1 explains the problem background and related studies. Chapter 2 introduces the technical stack and runtime environment. Chapter 3 derives requirements from the Nanpai scenario. Chapters 4 to 6 move from overall design to database design and detailed design. Chapter 7 presents the implemented functions, Chapter 8 reports the tests, and Chapter 9 summarizes the work together with the remaining limitations and future extensions.",
    205: "Keywords: enterprise informatization; resource-constrained manufacturers; micro-frontends; database-centric architecture; lightweight semantic enhancement",
    223: "现有国内研究已经给出了不少单点方法，但把中小制造企业的模块建设、流程联动、权限治理和轻量化语义增强放在同一平台中统筹落地的工程案例仍然偏少。缺口主要不在理论层，而在整合层：这些能力如何在资源有限的企业环境里同时工作，相关样本还不多。本文正是在这个位置继续展开。",
    230: "综合国内外研究现状可以看到，现有成果虽然已经覆盖数字化转型、流程建模和语义表达等多个方向，但真正落到中小制造企业核心业务系统时，缺口主要出现在整合层。很多研究按主题分别展开，业务模块、流程协同、权限治理和语义表达没有被放到同一平台中统筹考虑；不少成熟方案又更适用于大型企业或高标准化场景，对“小批量、多批次、频繁变更”的地方制造企业支撑不足。再往下走，问题就更具体了：一旦字段、表单或状态规则调整，依赖厚重后端服务层的传统系统会很快把维护成本抬高。",
    272: "本文所说的“本体”更接近面向业务表结构的语义增强，而不是另起一套独立知识图谱。它首先服务于现有系统：给表、字段和表间关系补充可解释的语义描述，让流程、权限和动态配置使用同一套对象口径。进一步看，这层语义已经进入 Agent 运行时。运行时不会让模型直接拼装接口路径，而是先判断当前对象属于数据表、流程状态还是语义关系，再落到白名单中的相应工具与接口，例如 `flash.data.table.ensure`、`flash.data.grid.list`、`flash.workflow.instance.start` 和 `flash.ontology.semantic.enrich`。",
    278: "本章涉及的技术并不是并列摆放的名词清单，而是围绕同一条运行链路分工协作：前端负责模块组织与交互承载，数据库负责数据结构和约束，PostgREST 缩短接口链路，BPMN 负责过程型业务，轻量化语义负责对象解释，Docker Compose 与 PM2 则把这些能力收束到一套可复现的工程环境中。对本文而言，重点不在“技术新不新”，而在“这套结构能不能在资源有限的企业场景里长期维护”。",
    295: "这一流程最突出的矛盾集中在信息衔接和状态透明度上。销售合同与销售订单之间缺少高效传递机制，客户、物料和数量等信息往往要重复录入；订单状态在销售、计划和仓库之间的传递也不够透明，销售人员很难及时知道“是否排产、是否完工、是否可发货”。系统因此需要支持从销售合同到销售订单的一键下推，并建立订单状态可视化机制。只有这样，销售、PMC 和仓库才可能围绕同一张订单形成一致判断。",
    359: "本章基于广东南派食品有限公司的调研资料，收拢了企业现状、角色边界、业务流程、功能需求和非功能需求。补料、质检延迟、客户物料映射、批次与保质期管理、纸质附件留存、冷库离线盘点以及细粒度权限控制等问题都被保留下来。它们并不是零散细节，而是后续设计必须回应的真实约束。接下来的总体设计与数据库设计，都围绕这些约束继续展开。",
    439: "库存流水表有两个直接用途。一个是给库存批次数量变化留下可以回查的来源记录。另一个是把静态库存状态和动态业务动作接起来，为后续追溯和业务分析保留过程数据。没有这张表，库存状态就只剩结果，没有过程。",
    538: "这里的关键不在于“能否写入一条入库记录”，而在于库存批次和库存流水不能分开处理。本文把两步都收进同一事务里，就是为了压掉几种最麻烦的中间状态：库存已经变了，流水没落库；或者流水写进去了，批次数量却没有同步更新。前者难查，后者更难查。",
    684: "结果为 23 项通过、0 项失败。现阶段测试说明，系统在登录鉴权、子应用访问、核心数据接口、流程定义读取、辅助运行能力以及核心进程在线状态等方面已具备稳定的基础可用性；围绕 Agent Runtime 的表格工具、本体关系查询工具和语义增强工具也完成了基础验证。也就是说，“智能体—数据表格应用—轻量化语义”这条链路已经能跑起来。移动端侧的冷库模式缓存、本地待处理记录和离线盘点入口同样完成了阶段性验证。",
}


def main():
    doc = Document(str(INPUT))
    changed = []
    for idx, text in REWRITE.items():
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].text = text
            changed.append(idx)

    doc.save(str(OUTPUT))
    doc.save(str(OUTPUT_EN))
    print(f"saved: {OUTPUT}")
    print(f"saved: {OUTPUT_EN}")
    print(f"changed: {len(changed)}")
    print(",".join(str(i) for i in changed))


if __name__ == "__main__":
    main()
