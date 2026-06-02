# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (c) 2026 林志荣

from docx import Document
from pathlib import Path


INPUT = Path(r"\\wsl.localhost\Ubuntu\home\lzr\eiscore\docs\tmp_thesis_v22_current.docx")
OUTPUT = Path(r"\\wsl.localhost\Ubuntu\home\lzr\eiscore\docs\毕业论文初稿v2.3_高风险压降版_2026-03-15.docx")
OUTPUT_EN = Path(r"\\wsl.localhost\Ubuntu\home\lzr\eiscore\docs\thesis_draft_v2_3_risk_reduced_2026-03-15.docx")


REWRITE = {
    197: "在南派食品的调研里，最先卡住人的不是少一个功能点，而是同一批业务记录总被拆开走。销售单、领料单、检验附件和库存变动并不待在一处：有的留在纸单上，有的在微信群里，有的又回到分散台账。只要其中一处补录慢了，后面追原料批次、仓位、补料和出库去向都会拖下来。对湛江不少中小制造企业来说，麻烦往往也是这样积出来的。预算不宽，运维人手又少，现场规则还会持续变化。",
    198: "所以 EISCore 先做底座，不先堆系统层级。核心数据、部分业务约束和访问边界先收回 PostgreSQL 与 PostgREST，再由微前端基座承载人事、物料、应用中心和移动端。数据库、接口、运行时与前端进程统一交给 Docker Compose 和 PM2 组织。先把人事治理、库存台账、应用配置、流程联动和语义增强几条主线接稳，后面的扩展才不至于越做越散。",
    231: "因此，本文继续追问的不是“还能补哪些单点功能”，而是怎样在资源有限的企业环境里，先把几条核心链路放到同一个底座里长期维护。",
    233: "围绕这个问题，本文把实现重点放在六个相互勾连的部分：人事与物料两条基础业务主线、以 PostgreSQL 和 PostgREST 为核心的数据库中心架构、基于 qiankun 的前端组织方式、应用中心中的闪念应用构建器、受控的 Agent Runtime，以及 BPMN 流程支撑与轻量化语义增强。它们不是独立卖点，而是围绕同一套业务对象和运行链路逐步落下来的。",
    272: "本文这里说的“本体”，不是要在系统旁边再建一套独立知识图谱。它先解决一个很实际的问题：同一张表、同一组字段到了流程、权限和动态配置里，解释口径不能一直漂。于是表、字段和表间关系先被补上语义说明。现在这层语义已经进入 Agent 运行时，但运行时并不允许模型自由拼接口。它会先判断当前对象落在数据表、流程状态还是语义关系，再调用白名单工具，例如 `flash.data.table.ensure`、`flash.workflow.instance.start` 或 `flash.ontology.semantic.enrich`。",
    295: "这一段流程里最难的并不是下单，而是状态总在不同岗位之间脱节。销售合同到销售订单之间缺少高效传递机制，客户、物料和数量等信息往往要重复录入；等订单往下走到销售、PMC 和仓库时，三边看到的状态又不完全一样。销售人员最难判断的恰恰是“是否排产、是否完工、是否可发货”。因此，系统既要支持从销售合同到销售订单的一键下推，也要把订单状态做成可视化，让相关岗位围着同一张订单说话。",
    434: "南派食品的仓储并不是一张平面表就能说明白。现场至少要区分仓库、库区和库位，盘点和定位都沿着这层空间关系走。`scm.warehouses` 因此通过 `parent_id` 组织成树形层级，后面不管是库存展示还是库位回查，都还沿用同一套结构。",
    439: "这张表不能少。批次数量为什么变了、变动对应哪张单据、发生在什么时间，都要从这里回查。否则库存台账只剩一个结果数字，追溯链很快就会断。",
    623: "在人事模块里，员工档案、组织结构和岗位管理已经形成可用页面。员工列表负责基础信息，组织结构页面展示部门层级，用户管理页面则把账号、岗位和角色关系收在一起。这样做有一个直接好处：后面做权限判断、流程候选人分派，甚至看某个账号为什么能看到某个菜单，都有统一入口可回查。",
    624: "部门、岗位和用户关系因此不再只是静态主数据。它们已经进入访问控制和流程分派这条链路。",
}


def main():
    doc = Document(str(INPUT))
    changed = []
    for idx, text in REWRITE.items():
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].text = text
            changed.append(idx)

    doc.save(str(OUTPUT))
    doc.save(str(OUTPUT_EN))
    print(f"saved: {OUTPUT}")
    print(f"saved: {OUTPUT_EN}")
    print(f"changed: {len(changed)}")
    print(",".join(str(i) for i in changed))


if __name__ == "__main__":
    main()
