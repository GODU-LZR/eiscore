# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (c) 2026 林志荣

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


BASE_DOC = Path(r"C:\Users\Twist\Desktop\论文\主稿\毕业论文初稿v2.8_格式目录终版_2026-04-07.docx")
OUT_DOC = Path(r"C:\Users\Twist\Desktop\论文\主稿\毕业论文初稿v2.8_数据库设计重写版_2026-04-12.docx")
FIG_ER = Path(r"\\wsl.localhost\Ubuntu\home\lzr\eiscore\docs\diagrams\图5-1_EISCore核心数据实体关系图.png")


TABLE_5_1 = [
    ["字段名", "类型", "长度/精度", "含义"],
    ["id", "integer", "-", "用户主键编号"],
    ["username", "text", "变长", "登录用户名"],
    ["password", "text", "变长", "登录密码"],
    ["role", "text", "变长", "默认系统角色标识"],
    ["avatar", "text", "变长", "用户头像地址"],
    ["permissions", "text[]", "数组", "用户附加权限集合"],
    ["full_name", "text", "变长", "用户姓名"],
    ["phone", "text", "变长", "联系电话"],
    ["email", "text", "变长", "邮箱地址"],
    ["dept_id", "uuid", "36位", "所属部门编号"],
    ["status", "text", "变长", "用户状态"],
    ["created_at", "timestamptz", "-", "创建时间"],
    ["updated_at", "timestamptz", "-", "更新时间"],
    ["position_id", "uuid", "36位", "所属岗位编号"],
]

TABLE_5_2 = [
    ["字段名", "类型", "长度/精度", "含义"],
    ["id", "uuid", "36位", "角色主键编号"],
    ["code", "text", "变长", "角色编码"],
    ["name", "text", "变长", "角色名称"],
    ["description", "text", "变长", "角色描述"],
    ["created_at", "timestamptz", "-", "创建时间"],
    ["updated_at", "timestamptz", "-", "更新时间"],
    ["sort", "integer", "-", "排序值"],
    ["dept_id", "uuid", "36位", "角色所属部门"],
]

TABLE_5_3 = [
    ["字段名", "类型", "长度/精度", "含义"],
    ["user_id", "integer", "-", "用户编号"],
    ["role_id", "uuid", "36位", "角色编号"],
    ["created_at", "timestamptz", "-", "关联建立时间"],
]

TABLE_5_4 = [
    ["字段名", "类型", "长度/精度", "含义"],
    ["id", "integer", "-", "物料主键编号"],
    ["batch_no", "text", "变长", "默认批次号或标识号"],
    ["name", "text", "变长", "物料名称"],
    ["category", "text", "变长", "物料分类"],
    ["weight_kg", "numeric", "10,2", "物料重量"],
    ["entry_date", "date", "-", "入档日期"],
    ["created_by", "text", "变长", "创建人"],
    ["properties", "jsonb", "-", "扩展属性"],
    ["version", "integer", "-", "版本号"],
    ["updated_at", "timestamp", "-", "更新时间"],
    ["dept_id", "uuid", "36位", "所属部门编号"],
]

TABLE_5_5 = [
    ["字段名", "类型", "长度/精度", "含义"],
    ["id", "uuid", "36位", "仓库节点主键"],
    ["code", "text", "变长", "仓库或库位编码"],
    ["name", "text", "变长", "仓库或库位名称"],
    ["parent_id", "uuid", "36位", "父节点编号"],
    ["level", "integer", "-", "层级，1为仓库、2为库区、3为库位"],
    ["sort", "integer", "-", "排序值"],
    ["status", "text", "变长", "启用状态"],
    ["manager_id", "integer", "-", "负责人编号"],
    ["capacity", "numeric", "15,2", "容量"],
    ["unit", "text", "变长", "容量单位"],
    ["properties", "jsonb", "-", "扩展属性，如布局、温湿度要求"],
    ["created_by", "text", "变长", "创建人"],
    ["created_at", "timestamptz", "-", "创建时间"],
    ["updated_at", "timestamptz", "-", "更新时间"],
    ["dept_id", "uuid", "36位", "所属部门编号"],
]

TABLE_5_6 = [
    ["字段名", "类型", "长度/精度", "含义"],
    ["id", "uuid", "36位", "批次记录主键"],
    ["material_id", "integer", "-", "物料编号"],
    ["batch_no", "text", "变长", "批次号"],
    ["warehouse_id", "uuid", "36位", "所在仓库节点编号"],
    ["available_qty", "numeric", "18,4", "可用数量"],
    ["locked_qty", "numeric", "18,4", "锁定数量"],
    ["unit", "text", "变长", "计量单位"],
    ["production_date", "date", "-", "生产日期"],
    ["expiry_date", "date", "-", "过期日期"],
    ["supplier", "text", "变长", "供应商"],
    ["purchase_price", "numeric", "15,4", "采购单价"],
    ["status", "text", "变长", "批次状态"],
    ["properties", "jsonb", "-", "扩展属性，如质检数据"],
    ["created_by", "text", "变长", "创建人"],
    ["created_at", "timestamptz", "-", "创建时间"],
    ["updated_at", "timestamptz", "-", "更新时间"],
    ["dept_id", "uuid", "36位", "所属部门编号"],
]

TABLE_5_7 = [
    ["字段名", "类型", "长度/精度", "含义"],
    ["id", "uuid", "36位", "流水主键"],
    ["transaction_no", "text", "变长", "单据编号"],
    ["transaction_type", "text", "变长", "事务类型，如入库、出库、调整"],
    ["material_id", "integer", "-", "物料编号"],
    ["batch_no", "text", "变长", "批次号冗余字段"],
    ["batch_id", "uuid", "36位", "批次记录编号"],
    ["warehouse_id", "uuid", "36位", "仓库节点编号"],
    ["quantity", "numeric", "18,4", "变动数量"],
    ["unit", "text", "变长", "计量单位"],
    ["before_qty", "numeric", "18,4", "变动前数量"],
    ["after_qty", "numeric", "18,4", "变动后数量"],
    ["related_doc_type", "text", "变长", "关联单据类型"],
    ["related_doc_no", "text", "变长", "关联单据编号"],
    ["transaction_date", "timestamptz", "-", "业务发生时间"],
    ["operator", "text", "变长", "操作人"],
    ["remark", "text", "变长", "备注"],
    ["approval_status", "text", "变长", "审批状态"],
    ["workflow_instance_id", "uuid", "36位", "关联流程实例编号"],
    ["properties", "jsonb", "-", "扩展属性"],
    ["created_by", "text", "变长", "创建人"],
    ["created_at", "timestamptz", "-", "创建时间"],
    ["dept_id", "uuid", "36位", "所属部门编号"],
]

TABLE_5_8 = [
    ["字段名", "类型", "长度/精度", "含义"],
    ["id", "uuid", "36位", "应用主键编号"],
    ["name", "varchar", "200", "应用名称"],
    ["description", "text", "变长", "应用描述"],
    ["category_id", "integer", "-", "所属分类编号"],
    ["app_type", "varchar", "20", "应用类型"],
    ["source_code", "jsonb", "-", "Flash应用源码"],
    ["config", "jsonb", "-", "数据应用配置"],
    ["bpmn_xml", "text", "变长", "流程定义XML"],
    ["icon", "varchar", "50", "图标标识"],
    ["status", "varchar", "20", "应用状态"],
    ["version", "varchar", "20", "应用版本"],
    ["created_by", "text", "变长", "创建人"],
    ["updated_by", "text", "变长", "更新人"],
    ["created_at", "timestamptz", "-", "创建时间"],
    ["updated_at", "timestamptz", "-", "更新时间"],
]

TABLE_5_9 = [
    ["字段名", "类型", "长度/精度", "含义"],
    ["id", "serial", "-", "流程实例主键"],
    ["definition_id", "int", "-", "流程定义编号"],
    ["business_key", "text", "变长", "业务主键标识"],
    ["current_task_id", "text", "变长", "当前任务节点编号"],
    ["status", "text", "变长", "流程实例状态"],
    ["variables", "jsonb", "-", "运行时变量集合"],
]


def find_paragraph(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise ValueError(f"Paragraph not found: {text}")


def remove_between(start_para, end_para):
    elem = start_para._p.getnext()
    while elem is not None and elem != end_para._p:
        nxt = elem.getnext()
        elem.getparent().remove(elem)
        elem = nxt


def insert_paragraph_before(end_para, text="", style=None, align=None):
    p = end_para.insert_paragraph_before(text, style=style)
    if align is not None:
        p.alignment = align
    return p


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold


def insert_table_before(doc, end_para, data):
    rows = len(data)
    cols = len(data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, row in enumerate(data):
        for c, value in enumerate(row):
            set_cell_text(table.cell(r, c), value, bold=(r == 0))
    end_para._p.addprevious(table._tbl)
    return table


def add_keep_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)


def insert_captioned_picture(doc, end_para, caption, img_path, width_in=5.7):
    pic_p = insert_paragraph_before(end_para, style="Body Text", align=WD_ALIGN_PARAGRAPH.CENTER)
    pic_p.add_run().add_picture(str(img_path), width=Inches(width_in))
    cap = insert_paragraph_before(end_para, caption, style="图片标题", align=WD_ALIGN_PARAGRAPH.CENTER)
    return pic_p, cap


def add_table_section(doc, end_para, subheading, intro, caption, explain, data):
    insert_paragraph_before(end_para, subheading, style="Heading 3")
    insert_paragraph_before(end_para, intro, style="Body Text")
    insert_paragraph_before(end_para, caption, style="表格标题", align=WD_ALIGN_PARAGRAPH.CENTER)
    insert_table_before(doc, end_para, data)
    insert_paragraph_before(end_para, explain, style="Body Text")


def main():
    doc = Document(str(BASE_DOC))
    start = find_paragraph(doc, "数据库设计")
    end = find_paragraph(doc, "系统详细设计")
    remove_between(start, end)

    # 5.1
    insert_paragraph_before(end, "数据库设计目标", style="Heading 2")
    insert_paragraph_before(
        end,
        "数据库设计是本设计的核心部分。EISCore 并没有把 PostgreSQL 仅作为被动存储层，而是围绕数据库中心架构，把权限边界、流程状态、批次库存、动态应用配置和语义描述都压到统一的数据底座上。因此，本章讨论的不只是“有哪些表”，而是这些表如何共同支撑系统运行、治理和扩展。",
        style="Body Text",
    )
    insert_paragraph_before(
        end,
        "结合南派食品的业务约束和当前系统实现范围，本设计的数据库目标主要有四点：其一，建立稳定的基础主数据模型，统一组织结构、用户、角色、物料、仓库等关键对象；其二，围绕库存批次、库存流水、入库出库和盘点链路组织仓储数据，为追溯与库存透明化提供基础；其三，为应用中心与工作流运行保留独立数据域，使动态应用、流程设计和审批协同能够在同一数据库内完成；其四，通过 RLS、JWT 声明和轻量语义描述结构，为后续的数据安全治理与语义增强留出稳定接口。",
        style="Body Text",
    )

    # 5.2
    insert_paragraph_before(end, "数据库总体结构设计", style="Heading 2")
    insert_paragraph_before(
        end,
        "EISCore 以 PostgreSQL 16 为核心，采用多 Schema 方式划分业务域。当前与本设计主线直接相关的 Schema 主要包括 public、hr、scm、app_center、workflow 和 app_data。与单一 Schema 下所有表混合存放的做法相比，多 Schema 更适合当前这种“基础主数据、仓储业务、流程运行和动态应用并存”的实现方式，也便于 PostgREST 暴露接口时按域控制访问边界。",
        style="Body Text",
    )
    insert_paragraph_before(
        end,
        "其中，public Schema 主要保存组织结构、用户、角色、权限、物料主数据及部分公共配置；hr Schema 主要承载员工档案、考勤、工资和人事扩展数据；scm Schema 主要承载仓库、库存批次、库存流水和盘点数据；app_center Schema 保存应用中心中的分类、应用、路由发布和流程状态映射；workflow Schema 保存流程定义、流程实例和任务分派；app_data Schema 则作为动态数据应用的承载域。这样的划分能够让“主数据、运行态、配置态”三类对象保持清晰边界，同时又能通过外键、业务键和状态映射形成统一数据底座。",
        style="Body Text",
    )

    # 5.3
    insert_paragraph_before(end, "核心数据实体关系设计", style="Heading 2")
    insert_paragraph_before(
        end,
        "从当前主线看，EISCore 的核心实体关系可以归纳为三条：组织与权限主线、物料与仓储主线、应用与流程主线。其核心实体关系如图5-1所示。",
        style="Body Text",
    )
    insert_captioned_picture(doc, end, "图5-1 EISCore核心数据实体关系图", FIG_ER, width_in=5.8)
    insert_paragraph_before(
        end,
        "在组织与权限主线中，departments、positions、users、roles 和 user_roles 构成了基础的组织授权关系。部门表支持层级结构，岗位表和用户表分别通过外键关联部门，角色表承载系统权限分类，用户与角色之间再通过关联表形成多对多关系。这条主线决定了用户能看到什么菜单、能处理哪些任务、能访问哪些数据域。",
        style="Body Text",
    )
    insert_paragraph_before(
        end,
        "在物料与仓储主线中，raw_materials、warehouses、inventory_batches 和 inventory_transactions 构成了库存业务链路。物料主数据表定义所有基础物料对象，仓库表通过树形结构描述仓库、库区和库位，库存批次表记录某个物料在某个位置上的批次库存，库存流水表则记录所有库存变化行为。通过这条主线，系统可以支撑采购入库、生产领料、补料出库、生产入库、销售出库以及盘点调整。",
        style="Body Text",
    )
    insert_paragraph_before(
        end,
        "在应用与流程主线中，categories、apps、workflow.definitions、workflow.instances、workflow.task_assignments、published_routes 和 workflow_state_mappings 构成了动态应用与流程协同能力的核心结构。应用表用于登记动态应用和流程应用，流程定义表与应用表建立关联，流程实例表和任务分派表承载运行态，而发布路由和状态映射表则分别用于前端挂载与业务状态写回。三条主线共同构成 EISCore 的数据库骨架。",
        style="Body Text",
    )

    # 5.4
    insert_paragraph_before(end, "组织与权限相关表设计", style="Heading 2")
    insert_paragraph_before(end, "部门表与岗位表", style="Heading 3")
    insert_paragraph_before(
        end,
        "部门表 public.departments 用于描述组织结构，主要字段包括 id、name、parent_id、leader_id、sort 和状态相关字段。通过 parent_id 自关联，系统可以表达树形组织结构，为人事归属和数据范围控制提供依据。岗位表 public.positions 则通过 dept_id 关联部门，描述部门内的岗位信息。",
        style="Body Text",
    )
    insert_paragraph_before(
        end,
        "这种设计的意义不只在于展示组织架构。对于南派食品当前场景，仓库人员、质检人员、销售文员和管理人员分属不同部门和岗位，部门与岗位信息会进一步参与角色配置、流程候选人筛选以及数据范围约束，因此组织表是权限治理的起点。",
        style="Body Text",
    )
    insert_paragraph_before(end, "用户表、角色表与用户角色关联表", style="Heading 3")
    insert_paragraph_before(
        end,
        "用户表 public.users 用于保存系统登录用户及其基本信息，关键字段包括 username、full_name、dept_id 和 position_id。角色表 public.roles 保存角色编码、角色名称以及所属部门等信息。由于同一用户在系统里可能承担审批、仓储、查看经营数据等多种职责，因此系统通过 public.user_roles 建立用户与角色之间的多对多关系。",
        style="Body Text",
    )
    insert_paragraph_before(
        end,
        "这一设计方式既能支持“一个用户一个主角色”的简单场景，也能支持“同一用户在不同业务中承担不同角色”的复杂场景。它与第六章中权限控制链路的实现保持一致：前端完成身份识别后，PostgREST 和 PostgreSQL 进一步根据角色声明和 RLS 策略裁剪数据访问范围。",
        style="Body Text",
    )

    # 5.5
    insert_paragraph_before(end, "物料与仓储相关表设计", style="Heading 2")
    insert_paragraph_before(end, "物料主数据表", style="Heading 3")
    insert_paragraph_before(
        end,
        "物料主数据是仓储链路的起点。当前系统中的 public.raw_materials 虽然命名为原料表，但在实现上承担了更广义的物料基础对象存储职责。除了名称、分类、部门归属等字段外，该表还通过 properties 字段兼容非标准扩展属性，用于适应食品加工场景中“同名物料、属性波动”的情况。",
        style="Body Text",
    )
    insert_paragraph_before(end, "仓库表与层级结构", style="Heading 3")
    insert_paragraph_before(
        end,
        "仓库表 scm.warehouses 采用树形结构设计，通过 parent_id 同时表示仓库、库区和库位三个层级。这一做法比把仓库、区域、库位拆成三张表更适合当前项目的规模，也便于移动端盘点和冷库模式复用同一套位置数据。对于食品企业常见的“一个仓库内再细分多个区域和位置”的管理方式，这种设计更贴近现场。",
        style="Body Text",
    )
    insert_paragraph_before(end, "库存批次表", style="Heading 3")
    insert_paragraph_before(
        end,
        "库存批次表 scm.inventory_batches 是当前数据库设计中最关键的业务表之一。该表以 material_id 和 warehouse_id 为基础，记录批次号、可用数量、锁定数量、生产日期、过期日期以及批次状态等信息。南派食品存在保质期管理、批次追溯和检验延迟等业务约束，因此这张表承担了库存台账和追溯链条中的关键角色。",
        style="Body Text",
    )
    insert_paragraph_before(end, "库存流水表", style="Heading 3")
    insert_paragraph_before(
        end,
        "库存流水表 scm.inventory_transactions 记录所有库存变化行为，包括采购入库、生产领料、补料出库、生产入库、销售出库和盘点调整等。它通过 material_id、batch_id 和 warehouse_id 与物料、批次和仓库建立关联，并保留交易类型、数量、关联单据和操作时间等信息。该表既是库存审计的依据，也是后续批次追溯、问题回放和状态对账的基础。",
        style="Body Text",
    )

    # 5.6
    insert_paragraph_before(end, "应用中心与流程相关表设计", style="Heading 2")
    insert_paragraph_before(end, "应用分类表与应用表", style="Heading 3")
    insert_paragraph_before(
        end,
        "应用分类表 app_center.categories 和应用表 app_center.apps 共同构成应用中心的基础结构。分类表负责区分数据应用、流程应用和其他应用类型；应用表保存应用名称、分类、类型、状态、配置以及 BPMN XML 等信息。由于 EISCore 并不只是固定页面集合，还需要支持动态应用配置、闪念应用草稿和流程设计，因此 app_center.apps 在数据库中扮演了“应用注册中心”的角色。",
        style="Body Text",
    )
    insert_paragraph_before(end, "发布路由与状态映射表", style="Heading 3")
    insert_paragraph_before(
        end,
        "发布路由表 app_center.published_routes 用于记录应用发布后的访问路径，便于前端基座动态挂载；状态映射表 app_center.workflow_state_mappings 用于把流程任务节点与具体业务表的状态字段建立关联，使流程推进时能够驱动业务状态写回。基座挂载页面时会读取前者，流程推进和写回业务状态时会读取后者，因此这两张表都直接参与系统运行。",
        style="Body Text",
    )
    insert_paragraph_before(end, "流程定义表、实例表与任务分派表", style="Heading 3")
    insert_paragraph_before(
        end,
        "流程定义表 workflow.definitions 用于保存 BPMN 流程定义，并通过 app_id 关联应用中心中的流程应用；流程实例表 workflow.instances 保存流程运行中的实例状态，包括当前任务节点、业务键和状态等；任务分派表 workflow.task_assignments 用于定义某个流程节点可由哪些角色或用户处理。通过这三类表，系统能够从静态流程定义延伸到动态流程运行与任务权限分派。",
        style="Body Text",
    )

    # 5.7
    insert_paragraph_before(end, "数据安全与权限控制设计", style="Heading 2")
    insert_paragraph_before(
        end,
        "本设计的数据安全控制并不依赖厚重的后端服务层，而是通过 PostgREST、JWT 声明和 PostgreSQL 的 RLS 策略共同落地。应用请求进入数据库前，PostgREST 会携带身份声明；数据库中的 RLS 策略再根据用户角色、部门范围和表级规则决定是否允许读取、写入或推进流程。对当前这种多子应用共享同一数据库的实现方式而言，这种“接口声明 + 数据库裁剪”的组合更容易保证边界一致。",
        style="Body Text",
    )
    insert_paragraph_before(
        end,
        "例如，应用中心中的 apps、published_routes、workflow_state_mappings 和 execution_logs 等表均配置了针对不同操作类型的访问限制；工作流相关表也通过角色声明约束写入和管理权限。这样做的直接好处，是让菜单可见性、按钮权限和数据写入边界不再各自为政，而是回到同一套数据库治理规则中。",
        style="Body Text",
    )

    # 5.8
    insert_paragraph_before(end, "语义增强相关设计", style="Heading 2")
    insert_paragraph_before(
        end,
        "轻量化语义增强并不是在数据库之外再搭一个独立知识库，而是在现有数据结构上增加一层可解释描述能力。当前设计主要围绕表级语义、列级语义和关系说明展开，用于支撑本体关系工作台、闪念应用对象解释以及数据应用字段理解。它的目标不是替代业务表，而是让应用中心、数据表格应用和受控 Agent 运行时能够更准确地理解“这张表是什么、字段代表什么、关系怎样组合”。",
        style="Body Text",
    )
    insert_paragraph_before(
        end,
        "从数据库设计角度看，语义增强的重点在于保留稳定的元数据锚点，而不是引入复杂推理引擎。也就是说，语义层必须依附现有的组织、物料、流程和应用表而存在，只有这样，后续在动态应用配置和对象解释中才能保持一致。",
        style="Body Text",
    )

    # 5.9 physical tables
    insert_paragraph_before(end, "关键物理表结构设计", style="Heading 2")
    insert_paragraph_before(
        end,
        "按照毕业设计说明书对数据库设计章节的要求，在说明总体结构和实体关系之后，还需要进一步给出关键物理表的字段结构。结合当前项目实现情况，本设计优先选取用户、角色、物料、库存、应用中心和流程运行等对主线影响最大的表进行展示。各表在本章中均先给出正文引入，再列出字段结构，避免只堆表格而缺少解释。",
        style="Body Text",
    )

    add_table_section(
        doc,
        end,
        "用户表",
        "为说明系统核心组织权限数据的基础结构，用户表的主要字段设计如表5-1所示。",
        "表5-1 用户表",
        "表5-1 展示了 public.users 的主要字段结构。该表承担账号信息、基础身份属性和组织归属信息的存储功能，是系统登录、身份识别和后续 RLS 数据裁剪的重要入口。",
        TABLE_5_1,
    )
    add_table_section(
        doc,
        end,
        "角色表",
        "为说明系统角色与权限分类的基础数据结构，角色表的主要字段设计如表5-2所示。",
        "表5-2 角色表",
        "表5-2 展示了 public.roles 的主要结构。该表用于定义系统角色编码、角色名称以及角色说明信息，并可与部门形成关联，为后续权限分配和流程节点候选角色配置提供基础。",
        TABLE_5_2,
    )
    add_table_section(
        doc,
        end,
        "用户角色关联表",
        "为说明用户与角色之间的关联方式，用户角色关联表的主要字段设计如表5-3所示。",
        "表5-3 用户角色关联表",
        "表5-3 展示了 public.user_roles 的字段结构。该表用于建立用户与角色之间的多对多关系，使系统能够支持一个用户对应多个角色，满足审批、仓储和管理查看等复合授权场景。",
        TABLE_5_3,
    )
    add_table_section(
        doc,
        end,
        "物料主数据表",
        "为说明物料主数据在系统中的组织方式，物料主数据表的主要字段设计如表5-4所示。",
        "表5-4 物料主数据表",
        "表5-4 展示了 public.raw_materials 的主要字段。该表虽然命名为原料表，但在当前系统实现中承担了较广义的物料基础信息存储职责，并通过 properties 字段兼容扩展属性。",
        TABLE_5_4,
    )
    add_table_section(
        doc,
        end,
        "仓库表",
        "为说明仓库基础信息的存储方式，仓库表的主要字段设计如表5-5所示。",
        "表5-5 仓库表",
        "表5-5 展示了 scm.warehouses 的主要结构。该表采用树形结构设计，可同时表示仓库、库区和库位三个层级，适合食品企业中分区管理和现场布局扩展的实际需要。",
        TABLE_5_5,
    )
    add_table_section(
        doc,
        end,
        "库存批次表",
        "为说明库存批次管理与保质期追踪的数据基础，库存批次表的主要字段设计如表5-6所示。",
        "表5-6 库存批次表",
        "表5-6 展示了 scm.inventory_batches 的字段结构。该表是本系统中最重要的仓储业务表之一，主要用于保存某种物料在某个仓位上的批次库存状态，为批次追溯、保质期管理和台账统计提供支撑。",
        TABLE_5_6,
    )
    add_table_section(
        doc,
        end,
        "库存流水表",
        "为说明库存收发流水的记录结构，库存流水表的主要字段设计如表5-7所示。",
        "表5-7 库存流水表",
        "表5-7 展示了 scm.inventory_transactions 的字段结构。该表记录所有出入库和库存变化行为，是连接库存状态与业务动作的核心表。系统通过该表保存单据号、业务类型、数量变化和关联单据信息，为库存审计和问题追溯提供依据。",
        TABLE_5_7,
    )
    add_table_section(
        doc,
        end,
        "应用注册表",
        "为说明应用中心中已注册应用的组织方式，应用注册表的主要字段设计如表5-8所示。",
        "表5-8 应用注册表",
        "表5-8 展示了 app_center.apps 的字段结构。该表用于保存应用中心中的应用定义信息，同时支持数据应用、流程应用、Flash 草稿应用和自定义应用，是动态应用管理能力的核心载体。",
        TABLE_5_8,
    )
    add_table_section(
        doc,
        end,
        "流程实例表",
        "为说明流程运行过程中的实例数据结构，流程实例表的主要字段设计如表5-9所示。",
        "表5-9 流程实例表",
        "表5-9 展示了 workflow.instances 的字段结构。该表用于保存流程运行态数据，是系统流程引擎与业务状态联动的核心表之一。与只保存设计稿的流程定义表不同，该表更关注具体业务对象当前运行到哪个任务节点以及实例状态如何变化。",
        TABLE_5_9,
    )

    # 5.10
    insert_paragraph_before(end, "本章小结", style="Heading 2")
    insert_paragraph_before(
        end,
        "本章围绕 EISCore 的数据库设计展开，依次说明了数据库设计目标、总体结构、核心实体关系、组织权限主线、物料仓储主线以及应用流程主线，并补充给出了九张关键物理表结构。整体来看，本系统数据库不仅承担数据存储职能，还承担权限控制、流程协同和语义增强等治理能力。这种数据库中心设计方式与本设计“轻量、可扩展、低运维成本”的总体目标保持一致，也为后续系统详细设计提供了稳定的数据基础。",
        style="Body Text",
    )

    doc.save(str(OUT_DOC))
    print(f"saved: {OUT_DOC}")


if __name__ == "__main__":
    main()
