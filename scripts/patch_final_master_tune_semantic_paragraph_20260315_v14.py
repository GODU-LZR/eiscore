# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (c) 2026 林志荣

from __future__ import annotations

from pathlib import Path

from docx import Document


SRC = Path(r"\\wsl.localhost\Ubuntu\home\lzr\eiscore\docs\毕业论文最终主稿_终版_补闪念语义冷库模式_润色版_2026-03-15.docx")
OUT = Path(r"\\wsl.localhost\Ubuntu\home\lzr\eiscore\docs\毕业论文最终主稿_终版_补闪念语义冷库模式_稳妥版_2026-03-15.docx")
OUT_EN = Path(r"\\wsl.localhost\Ubuntu\home\lzr\eiscore\docs\final_thesis_master_final_with_flash_semantic_coldmode_stable_2026-03-15.docx")


PARA_REPLACEMENTS = {
    272: "本文所说的“本体”更接近一种面向业务表结构的语义增强方式，而不是额外建设独立知识图谱。它的任务是给现有表、字段和表间关系补充可解释的语义描述，使流程、权限和动态配置在同一套数据结构上保持更一致的表达口径。当前项目已经把这部分语义描述接入 Agent 运行时，但方式并不是开放任意接口，而是先把能力整理成受控工具。例如，运行时中已经存在 flash.data.grid.list、flash.ontology.relation.list、flash.ontology.semantic.list 和 flash.ontology.semantic.enrich 等工具；模型先选择工具，再在既有权限范围内调用对应接口。这样处理以后，语义增强不再只是前端展示信息，也开始参与数据应用和辅助构建过程。",
}


def main() -> None:
    doc = Document(str(SRC))
    changed = 0
    for idx, text in PARA_REPLACEMENTS.items():
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].text = text
            changed += 1
    doc.save(str(OUT))
    doc.save(str(OUT_EN))
    print(f"saved: {OUT}")
    print(f"saved: {OUT_EN}")
    print(f"changed paragraphs: {changed}")


if __name__ == "__main__":
    main()
