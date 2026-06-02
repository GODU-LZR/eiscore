# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (c) 2026 林志荣

from __future__ import annotations

import re
from pathlib import Path

from docx import Document


DESKTOP_PAPER = Path(r"C:\Users\Twist\Desktop\论文")
SOURCE = DESKTOP_PAPER / "主稿" / "毕业论文初稿v2.8_模板格式终修版_2026-04-06.docx"
TARGET = DESKTOP_PAPER / "主稿" / "毕业论文初稿v2.8_降风险修订版_2026-04-06.docx"


def clean_caption_suffixes(text: str) -> str:
    text = re.sub(r"(图\d+-\d+[^\n。；]*?)\.(svg|png)\b", r"\1", text, flags=re.IGNORECASE)
    text = re.sub(r"(图\d+-\d+[^\n。；]*?)\.(SVG|PNG)\b", r"\1", text)
    text = re.sub(r"(图\d+-\d+[^\n。；]*?)\s*(svg|png)\b", r"\1", text, flags=re.IGNORECASE)
    return text


def rewrite(text: str) -> str:
    t = clean_caption_suffixes(text)

    replacements = [
        (
            "图1-1并不是目录翻版，而是把本设计的收敛路径压成一条清晰主线。图中从企业现实约束出发，先落到需求分析，再进入总体设计、关键技术路线、系统实现和测试验证，最后回到说明书成文。这样安排的目的，是把“南派食品到底遇到了什么问题”“EISCore 为什么这样设计”“后面各章分别承接哪一段工作”放在同一张图里讲清楚。 见图1-1。",
            "为把课题背景、需求来源、系统方案和后续章节之间的关系说清楚，本设计将整体研究路线整理为图1-1。图中从南派食品的现实约束出发，依次过渡到需求分析、系统设计、实现与验证几个阶段。这样处理的作用，是先把“问题从哪里来、设计为什么这样展开、后续各章分别承接什么内容”交代清楚。见图1-1。",
        ),
        (
            "此外，项目还引入了基于浏览器的 code-server 作为轻量 WebIDE 运行环境，用于支持部分开发与调试场景。这部分内容虽然偏工程实现，但对于说明系统整体运行环境和工具链仍然是必要的。",
            "此外，项目在开发过程中还使用了基于浏览器的 code-server 作为辅助开发环境。它不属于系统交付后的核心运行链路，但在多人协同调试、远程接入和临时修改场景中确实发挥了作用，因此在本节一并说明。",
        ),
        (
            "EISCore 采用分层架构设计，自上而下包括用户访问层、前端表现层、接口与运行层、数据与治理层以及支撑环境层，系统架构如图4-1所示。该图的重点不是列出所有页面，而是把“用户入口—前端子应用—接口运行层—数据库底座—部署环境”之间的架构关系压缩到同一张图中，使后续关键设计思路有统一参照。 见图4-1。",
            "EISCore 的整体架构如图4-1所示。图中把用户入口、前端子应用、接口运行层、数据库底座和部署环境放在同一张图里，目的是先交代系统各层之间怎样衔接，再为后续的关键设计思路提供统一参照。见图4-1。",
        ),
        (
            "流程启动与状态写回时序如图6-3所示。用户从业务页面或应用中心发起流程后，前端把流程定义编号、业务键和运行变量提交给 PostgREST，随后由 workflow.start_workflow_instance 创建流程实例并写入初始状态。当审批节点推进时，workflow.transition_workflow_instance 会结合 workflow.can_execute_task 和任务分派规则判断当前操作者是否有权限继续执行；若允许推进，则更新 workflow.instances、记录审批日志，并根据 app_center.workflow_state_mappings 将状态写回业务表。 见图6-3。",
            "图6-3展示了流程启动和状态写回的处理链路。当前端从业务页面或应用中心发起流程时，请求会先经由 PostgREST 进入数据库函数；`workflow.start_workflow_instance` 负责创建实例并写入初始状态，后续节点推进则由 `workflow.transition_workflow_instance` 接手。推进前，系统会先结合 `workflow.can_execute_task` 和任务分派规则校验当前操作者；校验通过后，再更新流程实例、记录审批日志，并依据 `app_center.workflow_state_mappings` 把结果写回对应业务表。见图6-3。",
        ),
        (
            "因此，本节详细设计要强调的是“语义怎样围绕现有系统落地”，而不是“理论上可以扩展出多复杂的本体系统”。对当前项目而言，轻量化语义的实际价值体现在帮助页面解释业务对象、帮助应用中心理解动态表结构，以及帮助受控工具识别当前对象类型。",
            "因此，本节关注的不是抽象本体如何继续扩展，而是语义能力怎样嵌入现有系统。对当前项目来说，轻量化语义主要承担三件事：给页面补充对象解释、给应用中心补充动态表结构说明、给受控工具提供对象类型判断依据。",
        ),
        (
            "图4-4展示了闪念应用构建页、草稿预览页、受控运行时、白名单工具、数据表格应用、流程对象和轻量语义之间的组织关系。",
            "图4-4展示了闪念应用构建页、草稿预览页、受控运行时、白名单工具、数据表格应用、流程对象和轻量语义之间的组织关系。这里强调的是它们在当前实现中的调用边界，而不是把闪念能力写成一个已经完全独立成熟的平台。",
        ),
        (
            "本设计围绕EISCore系统的设计与实现过程展开论述。第一章主要介绍课题的研究背景、研究意义以及国内外相关研究现状；第二章介绍系统所使用的关键技术与开发环境；第三章从系统目标、用户角色和业务场景出发，对系统需求进行分析；第四章给出系统总体架构、模块划分和关键设计思路；第五章重点描述数据库结构设计、核心数据模型以及本体语义增强机制设计；第六章说明系统中关键功能的详细设计与核心处理流程；第七章回到系统实现，结合主要页面和运行链路说明当前系统的实现结果；第八章对系统的主要测试情况进行总结；第九章总结全文工作，并说明当前系统边界和后续改进方向。",
            "本设计围绕 EISCore 的需求、设计、实现与验证展开。前文先交代课题背景、研究现状和系统所依赖的基础技术；随后进入需求分析、概要设计、数据库设计和详细设计；后文再回到系统实现、测试结果与当前边界。这样安排的目的，是让需求来源、设计取舍和实现结果能够前后对应，而不是把各章写成彼此割裂的说明片段。",
        ),
    ]

    for old, new in replacements:
        if old in t:
            t = t.replace(old, new)

    # English introduction: rewrite by markers
    if t.startswith("At present, digital transformation has become an important path"):
        t = (
            "For many small and medium-sized manufacturers in Zhanjiang, the real difficulty of digitalization is not whether software has been purchased, "
            "but whether records about orders, materials, approvals and warehouse movements can stay on one stable data chain. Nanpai Food is a typical case: "
            "once the same business record is split across paper sheets, chat messages and separate ledgers, later tracing and coordination become slow and fragile."
        )
    elif t.startswith("To address these issues, this design develops and implements EISCore"):
        t = (
            "To address this situation, this design develops EISCore as a lightweight enterprise information platform oriented to real business coordination. "
            "The system uses PostgreSQL as the core data carrier, exposes interfaces through PostgREST, organizes front-end modules with Vue 3 and qiankun, "
            "and further connects workflow support, lightweight semantic description, FlashBuilder, Agent Runtime and cold-storage PDA mode within the current implementation scope."
        )
    elif t.startswith("The design is organized as follows"):
        t = (
            "The remaining chapters move from background and technology foundation to requirements, overall design, database design, detailed design and system implementation. "
            "Testing results and current boundaries are then summarized at the end so that the design route and the implemented scope can be read in one continuous line."
        )

    return t


def main() -> None:
    doc = Document(str(SOURCE))
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            new_text = rewrite(para.text)
            if new_text != para.text:
                para.text = new_text

    # also clean table cells of .svg/.png suffixes if any
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text:
                        new_text = clean_caption_suffixes(para.text)
                        if new_text != para.text:
                            para.text = new_text

    TARGET.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(TARGET))
    print(TARGET)


if __name__ == "__main__":
    main()
