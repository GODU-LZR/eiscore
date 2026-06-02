# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (c) 2026 林志荣

from pathlib import Path
import shutil

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


DESKTOP = Path(r"C:\Users\Twist\Desktop\论文")
SOURCE = DESKTOP / "主稿" / "毕业论文初稿v2.9_英文前置说明对齐版_2026-04-14.docx"
OUT_THESIS = DESKTOP / "主稿" / "毕业论文初稿v2.9_英文前置说明精修版_2026-04-14.docx"
OUT_EN_DOCX = DESKTOP / "主稿" / "英文前置说明_单独版_2026-04-14.docx"
OUT_EN_MD = DESKTOP / "主稿" / "英文前置说明_单独版_2026-04-14.md"

TEMP_DIR = Path(r"C:\Temp\eiscore_thesis")
TEMP_DIR.mkdir(parents=True, exist_ok=True)
TEMP_SRC = TEMP_DIR / "v29_frontmatter_refine_src.docx"
TEMP_OUT = TEMP_DIR / "v29_frontmatter_refined.docx"
TEMP_EN = TEMP_DIR / "v29_frontmatter_only.docx"

REFINED_P3 = (
    "The specification is organized along the actual development path of EISCore rather than around "
    "isolated technical topics. It starts from the business setting and research basis, then moves "
    "to the technical foundation and requirement analysis, before turning to architecture, database "
    "structure, and detailed design. The later sections focus on implemented functions, operation "
    "results, and test verification, and the closing part summarizes the completed scope, current "
    "limits, and possible follow-up work."
)


def save_refined_thesis() -> tuple[str, str, str, str]:
    shutil.copyfile(SOURCE, TEMP_SRC)
    doc = Document(str(TEMP_SRC))

    intro_heading = None
    p1 = p2 = p3 = keywords = None
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text == "INTRODUCTION":
            intro_heading = idx
        elif intro_heading is not None and p1 is None and text:
            p1 = idx
        elif p1 is not None and p2 is None and text:
            p2 = idx
        elif p2 is not None and p3 is None and text:
            p3 = idx
        elif p3 is not None and text.startswith("KEYWORDS:"):
            keywords = idx
            break

    if None in (intro_heading, p1, p2, p3, keywords):
        raise RuntimeError("Could not locate English front matter paragraphs.")

    doc.paragraphs[p3].text = REFINED_P3
    doc.save(str(TEMP_OUT))
    shutil.copyfile(TEMP_OUT, OUT_THESIS)

    refined = Document(str(TEMP_OUT))
    p1_text = refined.paragraphs[p1].text.strip()
    p2_text = refined.paragraphs[p2].text.strip()
    p3_text = refined.paragraphs[p3].text.strip()
    keywords_text = refined.paragraphs[keywords].text.strip()
    return p1_text, p2_text, p3_text, keywords_text


def build_standalone_doc(p1: str, p2: str, p3: str, kw: str) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("INTRODUCTION")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    for text in (p1, p2, p3, kw):
        para = doc.add_paragraph(text)
        para.paragraph_format.first_line_indent = Pt(24)
        para.paragraph_format.space_after = Pt(6)

    doc.save(str(TEMP_EN))
    shutil.copyfile(TEMP_EN, OUT_EN_DOCX)

    md = "\n\n".join(["# INTRODUCTION", p1, p2, p3, kw]) + "\n"
    OUT_EN_MD.write_text(md, encoding="utf-8")


if __name__ == "__main__":
    p1, p2, p3, kw = save_refined_thesis()
    build_standalone_doc(p1, p2, p3, kw)
    print(str(OUT_THESIS))
    print(str(OUT_EN_DOCX))
    print(str(OUT_EN_MD))
