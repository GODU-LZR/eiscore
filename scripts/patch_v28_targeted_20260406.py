# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (c) 2026 林志荣

from pathlib import Path

from docx import Document


BASE = Path(r"C:\Users\Twist\Desktop\论文\主稿\毕业论文初稿v2.8_高概括段合规精修版_2026-04-06.docx")
DETAILED = Path(r"C:\Users\Twist\Desktop\论文\主稿\毕业论文初稿v2.8_详细设计链路修正版_2026-04-06.docx")
OUT = Path(r"C:\Users\Twist\Desktop\论文\主稿\毕业论文初稿v2.8_三处修复与乱码清理版_2026-04-06.docx")


def replace_para_text(paragraph, text: str) -> None:
    if paragraph.runs:
        first = paragraph.runs[0]
        first.text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def main() -> None:
    base = Document(str(BASE))
    detailed = Document(str(DETAILED))

    # Replace mojibake / broken detailed-design paragraphs with the clean,
    # chain-oriented wording from the detailed-chain revision.
    idxs = [
        349, 350, 352, 355, 357, 359, 363,
        413, 414, 416, 417, 421,
        424, 425, 427, 430, 433, 435, 436, 438,
    ]
    for idx in idxs:
        replace_para_text(base.paragraphs[idx], detailed.paragraphs[idx].text)

    # Chapter summary paragraph was still corrupted in both files; rewrite it.
    replace_para_text(
        base.paragraphs[440],
        "本章围绕数据库中心架构下的系统实现方式展开了详细设计，重点说明了权限控制、库存处理、流程审批与轻量化语义四条主线如何真正落到页面输入、接口调用、数据库函数执行以及状态写回链路之中。与概要设计相比，本章不再停留于模块职责和技术选型，而是进一步回答了 EISCore 在没有传统厚重后端分层的条件下，如何依靠 PostgREST、数据库函数、RLS 规则和状态映射表把关键业务处理稳定运行起来。",
    )

    # Tighten the English Introduction to reduce generic summary tone.
    replace_para_text(
        base.paragraphs[5],
        "For resource-constrained manufacturers in Zhanjiang, the hard part is not the first software rollout. The harder part is keeping order records, stock movements, inspection attachments, and approval states on one traceable data path. At Nanpai Food, once these records are split across paper forms, chat messages, and separate ledgers, tracing a batch from receipt to supplement issue and final outbound delivery quickly becomes slow, fragmented, and error-prone.",
    )
    replace_para_text(
        base.paragraphs[6],
        "EISCore addresses this situation with a database-centric structure rather than a heavy middle-tier design. PostgreSQL carries core schemas, selected business constraints, and database functions; PostgREST exposes these capabilities as uniform endpoints; Vue 3 with qiankun organizes the base application, human-resources pages, material pages, the application center, and mobile entry points. Workflow runtime, semantic descriptions, FlashBuilder, Agent Runtime, and the cold-storage mode are all attached to this same execution chain, so the system can keep critical records on one stable path instead of scattering them across isolated tools.",
    )

    # Reduce template-like overview wording in chapter 1 / 4 / 6.
    replace_para_text(
        base.paragraphs[30],
        "结合前述问题分析，本设计的研究内容不按抽象概念展开，而是围绕 EISCore 已经落地或已形成原型链路的部分来组织，重点包括数据库中心架构、微前端组织方式、库存与流程协同、轻量化语义增强以及移动端冷库模式等内容。",
    )
    replace_para_text(
        base.paragraphs[32],
        "从当前实现情况看，系统已经形成可运行的核心链路，但覆盖范围仍然保持克制。人事、物料、应用中心和移动端等模块已经能够支撑真实页面操作，流程设计与语义增强也完成了阶段性落地，因此本设计讨论的不只是方案设想，而是建立在实际实现与联调基础上的工程总结。",
    )
    replace_para_text(
        base.paragraphs[61],
        "传统前后端分离系统通常要额外维护后端接口层，再由服务端把数据库操作包装成 API。本设计没有沿用这条路线，而是让 PostgREST 直接基于数据库模式、函数和权限规则暴露统一接口。这样处理并不是为了省略设计，而是为了把开发重心从重复接口封装转回到表结构、数据库函数、RLS 与状态写回这些真正影响系统稳定性的环节上。",
    )

    base.save(str(OUT))


if __name__ == "__main__":
    main()
