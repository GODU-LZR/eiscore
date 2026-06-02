# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (c) 2026 林志荣

from __future__ import annotations

from collections import defaultdict
from pathlib import Path
import re

from docx import Document


SRC = Path(r"C:\Users\Twist\Desktop\论文\主稿\毕业论文初稿v2.8_降风险与引用修订版_2026-04-06.docx")
TARGET = Path(r"C:\Users\Twist\Desktop\论文\主稿\毕业论文初稿v2.8_降风险与引用终修版_2026-04-06.docx")
AUDIT_DIR = Path(r"C:\Users\Twist\Desktop\论文\清单说明\v28_降风险与引用终修版_图表引用核查_2026-04-06")


def text(p):
    return p.text.strip().replace("\n", " ")


def replace_first(doc: Document, needle: str, new_text: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if needle in text(p):
            p.text = new_text
            return i
    raise ValueError(f"not found: {needle}")


def audit_refs(doc: Document) -> str:
    fig_ids: dict[str, list[int]] = defaultdict(list)
    tab_ids: dict[str, list[int]] = defaultdict(list)
    body_refs: dict[str, list[int]] = defaultdict(list)

    fig_pat = re.compile(r"图\d+-\d+")
    tab_pat = re.compile(r"表\d+-\d+")

    for i, p in enumerate(doc.paragraphs):
        t = text(p)
        if not t:
            continue
        is_caption = False
        if t.startswith("图"):
            ids = fig_pat.findall(t)
            if ids:
                is_caption = True
                for fid in ids:
                    fig_ids[fid].append(i)
        if t.startswith("表"):
            ids = tab_pat.findall(t)
            if ids:
                is_caption = True
                for tid in ids:
                    tab_ids[tid].append(i)
        if not is_caption:
            for fid in fig_pat.findall(t):
                body_refs[fid].append(i)
            for tid in tab_pat.findall(t):
                body_refs[tid].append(i)

    lines: list[str] = []
    lines.append("# 图表正文引用核查")
    lines.append("")
    lines.append(f"- 输入文件：`{TARGET}`")
    lines.append("")
    lines.append("## 图引用情况")
    lines.append("")
    for fid in sorted(fig_ids.keys(), key=lambda s: tuple(map(int, s[1:].split("-")))):
        refs = body_refs.get(fid, [])
        status = "已引用" if refs else "缺正文引用"
        lines.append(f"- `{fid}`：{status}；图题段落 {fig_ids[fid]}；正文引用段落 {refs}")
    lines.append("")
    lines.append("## 表引用情况")
    lines.append("")
    for tid in sorted(tab_ids.keys(), key=lambda s: tuple(map(int, s[1:].split("-")))):
        refs = body_refs.get(tid, [])
        status = "已引用" if refs else "缺正文引用"
        lines.append(f"- `{tid}`：{status}；表题段落 {tab_ids[tid]}；正文引用段落 {refs}")
    lines.append("")
    missing = [k for k in fig_ids if k not in body_refs] + [k for k in tab_ids if k not in body_refs]
    lines.append("## 结论")
    lines.append("")
    if missing:
        lines.append(f"- 仍缺正文引用：{', '.join(missing)}")
    else:
        lines.append("- 当前识别到的图表编号在正文中均有引用。")
    lines.append("")
    return "\n".join(lines)


doc = Document(str(SRC))
changed: list[tuple[int, str]] = []

replacements = [
    (
        "图3-6给出的不是菜单罗列",
        "图3-6用于说明本设计当前已经落地的核心能力如何按域组织。其中，人事管理、物料与仓储、应用中心和移动端构成直接面向用户的业务入口；主数据、流程与权限支撑负责维持这些入口之间的数据一致性和状态联动。这样处理的目的，不是把功能名称重新列一遍，而是说明后续概要设计为何围绕这些能力展开。",
    ),
    (
        "需求分析走到这里，已经可以看出本系统后续为何要采用数据库中心架构",
        "从上述需求可以看出，南派食品真正缺的并不是若干孤立页面，而是一套能够把订单、库存、流程状态和追溯信息稳定串起来的底座。也正因为如此，后续概要设计才会采用数据库中心架构、流程状态映射和轻量语义支撑，把经常变化的业务字段、流程节点和数据解释统一放到可维护的结构中。",
    ),
    (
        "在当前实现中，闪念应用构建页负责组织草稿输入和目标对象",
        "图4-4说明的是一条已经在当前系统中打通的联动链路。具体做法是：用户先在 FlashBuilder 页面输入草稿目标和页面意图，再由受控运行时根据对象类型挑选允许调用的工具；语义结果用于补充对象解释，数据表格应用负责承接字段、表单和页面草稿，最后把结果回写到应用中心和相关 Schema 中。这样写的重点是说明系统如何协同，而不是把几个名词并排放在同一句里。",
    ),
    (
        "把这条链路进一步落到当前实现，可以更具体地描述为：用户首先在基座登录页完成人员身份识别",
        "把权限链路落到当前实现，可以按三个连续步骤来看。第一步，用户在基座登录页完成身份识别，登录结果写入前端状态对象，并交给路由守卫和按钮控制使用。第二步，请求进入 PostgREST 时携带 JWT 声明，接口层并不自己维护一套新的权限规则，而是把身份信息继续传给数据库。第三步，数据库函数和 RLS 再根据当前用户、目标表和目标操作决定是否允许读取、写入或推进流程。这样描述之后，图6-4才真正对应到“权限怎样跑起来”这一实现问题。",
    ),
    (
        "本设计将其拆解为“审批页面或流程配置页面、PostgREST 接口、流程函数、流程实例表、审批记录表和状态映射表”六类对象。",
        "为避免把流程能力写成抽象概念，这里将其拆成六类实际对象：审批页面或流程配置页面、PostgREST 接口、流程函数、流程实例表、审批记录表和状态映射表。这样拆分的目的，是让读者直接看到页面操作怎样进入接口、接口怎样触发流程函数、函数又怎样更新实例表和业务状态，而不是只知道系统“支持审批”。",
    ),
    (
        "第一类是表级语义，用来说明某张表在业务上表示什么对象",
        "轻量语义模块在当前实现里主要做三件事。其一，给表补业务解释，例如某张表对应什么对象、归属哪个模块、承担什么职责。其二，给字段补语义说明，例如字段代表什么含义、由谁维护、会影响哪些业务规则。其三，把这些结果继续提供给应用中心、FlashBuilder 和受控运行时使用，使页面草稿、字段配置和数据解释能够引用同一套语义结果。",
    ),
    (
        "第三，应用中心、FlashBuilder 与 Agent Runtime 在当前实现中怎样使用这些语义结果。",
        "第三部分关注这些语义结果在当前系统里怎样真正被用起来。应用中心会读取语义结果去辅助数据应用配置；FlashBuilder 会结合对象说明和字段语义生成草稿；受控运行时再按对象类型和允许工具继续处理这些结果。换句话说，语义模块并不是孤立展示层，而是已经参与到页面配置和受控生成链路中。",
    ),
    (
        "用户可在应用中心中查看已有应用、进入数据应用配置页面，并对相关数据应用进行配置和管理。",
        "当前应用中心已经支持查看已有应用、进入数据应用配置页面，并对字段、表单和页面草稿进行管理。这里的实现重点不是单纯列出一个应用列表，而是说明系统已经具备继续组织数据应用的能力：对象先在应用中心注册，再通过配置页面补充字段、布局和发布信息，后续草稿和语义结果也能够落到这一入口中。",
    ),
    (
        "图6-4 权限控制时序图",
        "图6-4 权限控制时序图",
    ),
]

for needle, new_text in replacements:
    idx = replace_first(doc, needle, new_text)
    changed.append((idx, needle[:24]))

# 单独处理图6-4相关长段，避免与图题连成一段
for i, p in enumerate(doc.paragraphs):
    t = text(p)
    if t.startswith("图6-4 权限控制时序图") and "本节小结" in t:
        p.text = (
            "图6-4给出了权限控制链路的关键时序。本节讨论的重点，不是权限表怎样建，而是从页面登录、路由守卫、"
            "PostgREST 传递身份，到数据库函数与 RLS 最终裁决访问资格这一整条链路怎样在当前实现中真正跑通。"
            "通过这张时序图，可以更清楚地看到没有传统后端 controller/service 时，权限控制究竟落在前端状态、接口声明还是数据库边界。"
        )
        changed.append((i, "图6-4长段"))
        break

TARGET.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(TARGET))

AUDIT_DIR.mkdir(parents=True, exist_ok=True)
(AUDIT_DIR / "audit_refs.md").write_text(audit_refs(Document(str(TARGET))), encoding="utf-8")
(AUDIT_DIR / "changes.md").write_text(
    "# 本轮定点修改\n\n" + "\n".join(f"- 段落 {idx}: {label}" for idx, label in changed) + "\n",
    encoding="utf-8",
)
