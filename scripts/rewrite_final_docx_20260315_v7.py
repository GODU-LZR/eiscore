from __future__ import annotations

from pathlib import Path

from docx import Document


SRC = Path("/home/lzr/eiscore/docs/毕业论文总稿_按校模板终稿_修订版6_2026-03-15.docx")
OUT = Path("/home/lzr/eiscore/docs/毕业论文总稿_按校模板终稿_修订版7_2026-03-15.docx")


PARA_REPLACEMENTS = {
    202: (
        "At present, digital transformation has become an important path for manufacturing enterprises to improve operational efficiency and market competitiveness. However, for many small and medium-sized manufacturing enterprises in Zhanjiang, information system construction still faces practical constraints such as limited financial investment, insufficient technical support, difficult customization and high maintenance costs. Guangdong Nanpai Food Co., Ltd., which is used as the practical background of this study, has obvious business characteristics such as multiple raw material batches, non-standard material attributes and frequent process changes. Under such conditions, problems such as scattered information, weak traceability, unclear permission boundaries and low collaboration efficiency are likely to occur in material management, inventory records, personnel governance and workflow coordination. Traditional large-scale ERP systems are often expensive and difficult to deploy quickly, while lightweight SaaS tools are usually limited in customization and long-term expansion. Therefore, it is of practical significance to design a lightweight, extensible and low-maintenance enterprise information system core platform for small and medium-sized manufacturing enterprises.",
        "For small and medium-sized manufacturing enterprises in Zhanjiang, digitalization is no longer only a question of whether software is available. A more practical issue is whether orders, inventory, personnel records, workflow states and quality information can be connected into a continuous data chain. Guangdong Nanpai Food Co., Ltd., which serves as the business background of this study, shows several typical difficulties of this kind of enterprise: raw materials circulate in multiple batches, material attributes are not fully standardized, process changes occur frequently, and many records still rely on manual transfer. Under these conditions, scattered information, weak traceability, unclear permission boundaries and low collaboration efficiency become common problems. Large ERP packages are often too costly and heavy for such enterprises, while lightweight tools are easier to start with but usually weaker in customization and long-term extension. This is the practical context in which a lightweight and maintainable enterprise information system core platform becomes necessary."
    ),
    203: (
        "To address these issues, this thesis designs and implements EISCore, an enterprise information system core platform oriented to small and medium-sized manufacturing enterprises in Zhanjiang. The system adopts a database-centric architecture, with PostgreSQL serving as the core carrier of business data and part of the business logic, while PostgREST is used to automatically expose data interfaces and reduce repetitive backend interface development. On the frontend side, the system is built with Vue 3 and organized through a micro-frontend architecture based on qiankun, so that modules such as human resource management, material management, application center and mobile services can run collaboratively on a unified platform. In terms of business implementation, the system focuses on two core domains, namely HMS and MMS, and has completed the design and implementation of employee archives, organizational structure, user and role management, attendance management, material master data, warehouse management, inventory ledger, stock-in, stock-out and stock checking functions. In addition, BPMN-based workflow support and a lightweight semantic enhancement mechanism are introduced to improve workflow coordination, data expression and future extensibility. Overall, this study is concerned not only with whether the system can be implemented, but also with whether it can remain sustainable in business expression, structural extension and further intelligent evolution.",
        "Based on these problems, this thesis designs and implements EISCore as a lightweight information system core for small and medium-sized manufacturing enterprises. The system follows a database-centric structure: PostgreSQL carries the core data structures and part of the business constraints, while PostgREST exposes data interfaces and reduces repetitive backend API work. On the frontend side, Vue 3 and qiankun are used to organize a unified entry with several coordinated modules, including human resource management, material management, the application center and mobile services. In the current implementation scope, the system has already formed two stable business lines around personnel governance and material circulation, covering employee archives, organizational structure, user-role management, attendance, material master data, warehouse management, inventory ledger, stock-in, stock-out and stock checking. BPMN-based workflow support and a lightweight semantic enhancement mechanism are further added to connect workflow states, permission control and future extension needs. Therefore, the study focuses not only on whether the system can run, but also on whether it can remain maintainable and extensible in a real enterprise environment."
    ),
    204: (
        "This thesis is organized as follows. Chapter 1 introduces the research background, research significance, related studies at home and abroad, and the main contents of this thesis. Chapter 2 presents the key technologies and development environment used in the system. Chapter 3 analyzes the system requirements from the perspectives of business objectives, user roles and application scenarios. Chapter 4 describes the overall system design, including architecture, module division and major design ideas. Chapter 5 focuses on database design, core data structures and the lightweight semantic enhancement mechanism. Chapter 6 gives the detailed design of key modules, including user and permission control, inventory processing, workflow linkage and semantic support. Chapter 7 presents the implementation results of the system. Chapter 8 verifies the system through testing and result analysis. Finally, Chapter 9 summarizes the completed work, points out the current limitations of the system and discusses future improvements.",
        "The thesis is arranged in nine chapters. Chapter 1 explains the research background, related studies and the research focus of this thesis. Chapter 2 introduces the main technologies and the runtime environment adopted in EISCore. Chapter 3 derives the system requirements from the enterprise scenario, business roles and operational problems. Chapter 4 presents the overall design of the system, and Chapter 5 explains the database structure together with the semantic enhancement design. Chapter 6 further describes the detailed design of key parts such as permission control, inventory processing, workflow linkage and semantic support. Chapter 7 shows the implementation results, Chapter 8 reports the testing work, and Chapter 9 concludes the thesis and discusses the remaining limitations together with possible future improvements."
    ),
    228: (
        "总体而言，国外研究在低代码、微前端和数据库中心架构方面已经积累了较成熟的经验，但多数研究仍聚焦单项技术或单一组织环境。对于本文所面对的“中小制造企业核心业务系统”场景，尤其是流程、权限、动态扩展和语义建模并存的情况，如何形成一套兼顾轻量部署和持续扩展的底座方案，仍有继续整合和落地的空间。",
        "从国外研究的推进路径来看，低代码、微前端和数据库中心架构都已形成较成熟的方法积累，但这些成果多数仍围绕单项技术展开。具体到本文面对的中小制造企业核心业务系统场景，问题不只是选哪一种技术，而是如何把流程、权限、动态扩展和语义表达放进同一套底座结构中协同落地。现有研究在这一层面的整合性案例仍然有限，这也构成了本文继续展开设计与实现的现实依据。"
    ),
    278: (
        "本章并不是把相关技术逐一做教材式介绍，而是说明这些技术在 EISCore 中分别承担什么职责。Vue 3、Vite 与 qiankun 负责组织前端多模块运行，PostgreSQL、PostgREST 与 RLS 负责数据、接口和权限边界，BPMN 负责过程型业务建模，轻量化语义机制负责增强数据结构的可解释性，Docker Compose 与 PM2 则共同支撑系统运行环境。这些技术并非彼此孤立，而是围绕“中小制造企业信息化底座”这一目标被组合起来。",
        "本章所讨论的技术，并不是彼此独立堆叠在一起的工具集合，而是围绕 EISCore 的实际运行链路分工协作。前端侧由 Vue 3、Vite 和 qiankun 负责统一入口与多模块组织，数据库侧由 PostgreSQL、PostgREST 和 RLS 负责数据结构、接口访问和权限边界，流程侧借助 BPMN 支撑过程型业务，语义侧则通过轻量化机制增强表和字段的可解释性，部署侧再由 Docker Compose 与 PM2 维持运行环境的一致性。这样组合这些技术，目的不是展示技术种类，而是把中小制造企业最需要的核心链路稳定下来。"
    ),
    722: (
        "本章围绕系统测试目标、测试环境、测试方法和测试结果，对 EISCore 当前阶段的运行效果进行了验证。测试结果表明，系统在登录鉴权、页面访问、核心数据接口、流程支撑能力和运行状态等方面已经通过阶段性测试，论文中实现的核心功能具备较好的可运行性。与此同时，自动化测试、异常场景覆盖和性能评估仍有补充空间，这些内容也是后续优化的重点。",
        "本章测试工作的重点，在于验证 EISCore 当前已经完成的主干链路是否能够稳定运行。现有结果表明，登录鉴权、页面访问、核心数据接口、流程支撑和运行状态等关键环节已经完成阶段性验证，论文中实现的核心功能具备基本可运行性。与此同时，自动化测试覆盖、异常场景处理和性能评估仍需继续补强，因此后续优化工作应更多放在持续验证能力的完善上。"
    ),
}


def main() -> None:
    doc = Document(str(SRC))
    changed = 0
    for idx, (_, new_text) in PARA_REPLACEMENTS.items():
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].text = new_text
            changed += 1
    doc.save(str(OUT))
    print(f"saved: {OUT}")
    print(f"changed paragraphs: {changed}")


if __name__ == "__main__":
    main()
