from docx import Document
from pathlib import Path


INPUT = Path(r"\\wsl.localhost\Ubuntu\home\lzr\eiscore\docs\tmp_thesis_v24_current.docx")
OUTPUT = Path(r"\\wsl.localhost\Ubuntu\home\lzr\eiscore\docs\毕业论文初稿v2.5_问题修正版_2026-03-15.docx")
OUTPUT_EN = Path(r"\\wsl.localhost\Ubuntu\home\lzr\eiscore\docs\thesis_draft_v2_5_issue_fixed_2026-03-15.docx")


REWRITE = {
    197: "南派食品这类企业在信息化上最容易出问题的地方，不是少装一套软件，而是同一批业务记录总在不同介质之间来回转手。库存台账、领料记录、检验附件和审批状态一旦分散在纸单、聊天消息和独立台账里，后面再追一批原料进了哪个仓位、有没有发生补料、最后跟到哪张出库单，速度就会明显慢下来。湛江不少中小制造企业面对的现实约束也差不多：预算有限，系统定制成本高，运维人手不足，现场规则又经常变化。",
    198: "本文据此设计并实现 EISCore，希望先把这些容易断开的链路收回到同一底座里。数据库、接口与运行时服务通过 Docker Compose 统一编排，基座、人事、物料、应用中心和移动端等模块在前端侧协同运行。这样的组织方式并不追求把系统做得更重，重点还是先稳住人事治理、物料流转、应用配置、流程联动和语义增强这几条主线，再为后续扩展留下余地。",
    199: "全文围绕 EISCore 的设计、实现与验证展开。前面的章节先交代课题背景、相关研究、企业场景和需求约束，随后进入总体设计、数据库设计与详细设计；后面的章节再回到系统实现、测试结果和当前边界，说明这套底座已经把哪些链路落稳，哪些部分仍保留为后续扩展方向。",
    202: "At Nanpai Food, the practical problem is not the simple presence or absence of software. Trouble appears when one business record is split across paper forms, chat messages and separate ledgers. Stock movements, inspection files and workflow states no longer stay on the same path for long, and traceability becomes slower as soon as one step falls behind.",
    203: "This thesis therefore develops EISCore as a lightweight information core for resource-constrained manufacturers. Instead of starting from a heavier service stack, the project keeps core data, part of the business constraints and access boundaries close to PostgreSQL, while PostgREST shortens routine interface work. On the client side, Vue 3 and qiankun keep one entry and separate the modules that have already become stable: human resources, material management, the application center and mobile services.",
    204: "Within the current scope, the system has already stabilized two main lines around personnel and materials. Workflow support, FlashBuilder, controlled agent tools, semantic enhancement and the cold-storage PDA mode are added around the same runtime chain. The point is not to show a long list of technologies. The point is to keep a few core links maintainable under the actual conditions of Nanpai Food.",
    205: "Keywords: manufacturing SMEs; information core; micro-frontends; database-centered design; semantic support",
    268: "在本系统中，BPMN 更像一层过程支撑，而不是单独摆着看的流程图。它把原本分散在页面和状态字段里的过程型业务收回来，让审批节点、流转记录和业务状态能接在一起，后面的回写与查询也更容易落到同一条链路上。",
    272: "本文这里所说的“本体”，并不是要在系统外再搭一套独立知识图谱。它先处理一个更具体的问题：同一张表、同一组字段到了流程、权限和动态表单里，解释口径经常会飘。于是项目先给现有对象补上语义说明。现在这层语义已经进入 Agent 运行时，不过运行时并不会让模型自己拼接口，而是先识别对象类型，再转到允许调用的工具与接口。",
    377: "在接口与运行层，系统主要由 PostgREST、工作流运行组件和辅助运行组件构成。PostgREST 负责把数据库中的表、视图和函数映射成接口；工作流运行组件承载流程实例和任务流转；辅助运行组件则接住闪念应用相关的 Agent 接口、草稿同步、附件上传和工具编排。运行时目前采用白名单工具组织方式，重点是把智能体调用收在既有权限边界内。这样一来，数据表格应用、流程状态和轻量化语义可以继续联动，但不会变成任意接口自由拼接。",
    642: "应用中心模块当前已完成首页、数据应用配置、应用运行页面以及闪念应用构建器。用户可以查看已有应用、维护字段与配置、查看运行效果，也可以通过 FlashBuilder 进入对话式草稿构建流程。当前这条链路已经具备草稿读取与保存、附件上传、预览、发布路由写入以及与 Agent Runtime 协同处理的基础能力。运行时提供的是受控工具，因此闪念应用可以继续联动数据表格应用和语义增强结构，但仍保持在既有权限边界内。"
}


def main():
    doc = Document(str(INPUT))
    changed = []
    for idx, text in REWRITE.items():
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].text = text
            changed.append(idx)

    doc.save(str(OUTPUT))
    doc.save(str(OUTPUT_EN))
    print(f"saved: {OUTPUT}")
    print(f"saved: {OUTPUT_EN}")
    print(f"changed: {len(changed)}")
    print(",".join(str(i) for i in changed))


if __name__ == "__main__":
    main()
