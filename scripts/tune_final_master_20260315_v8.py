from __future__ import annotations

from pathlib import Path

from docx import Document


SRC = Path("/home/lzr/eiscore/docs/毕业论文最终主稿_2026-03-15.docx")
OUT = Path("/home/lzr/eiscore/docs/毕业论文最终主稿_稳妥版_2026-03-15.docx")
OUT_EN = Path("/home/lzr/eiscore/docs/final_thesis_master_stable_2026-03-15.docx")


PARA_REPLACEMENTS = {
    202: "For small and medium-sized manufacturing enterprises in Zhanjiang, the key issue of digitalization is not simply whether a software system is available, but whether orders, inventory, personnel records, workflow states and quality information can be connected into one usable data chain. Guangdong Nanpai Food Co., Ltd., which provides the business background of this study, reflects several typical problems in this kind of environment: raw materials circulate in multiple batches, material attributes vary from batch to batch, process changes are frequent, and many records still depend on manual transfer. Under these conditions, information becomes fragmented, traceability weakens, permission boundaries are blurred, and collaboration efficiency declines. Large ERP systems are often too heavy and costly for this scenario, while lightweight tools are easier to adopt but usually weaker in customization and later extension. This is the practical setting in which a lighter and more maintainable system core is needed.",
    203: "To respond to these problems, this thesis designs and implements EISCore as a lightweight enterprise information system core for small and medium-sized manufacturing enterprises. The system uses a database-centric structure in which PostgreSQL carries the core data model together with part of the business constraints, while PostgREST provides interface exposure and reduces repetitive backend API work. On the frontend side, Vue 3 and qiankun are used to organize a unified entry with coordinated modules such as human resource management, material management, the application center and mobile services. Within the current implementation scope, the project has already formed two relatively stable business lines around personnel governance and material circulation, covering employee archives, organizational structure, user-role management, attendance, material master data, warehouse management, inventory ledger, stock-in, stock-out and stock checking. BPMN-based workflow support and a lightweight semantic enhancement mechanism are also introduced so that workflow states, permission control and later extension needs can be connected more consistently.",
    204: "The thesis is divided into nine chapters. Chapter 1 explains the research background, related studies and the research focus of this work. Chapter 2 introduces the main technologies and the runtime environment used in EISCore. Chapter 3 derives system requirements from the enterprise scenario, business roles and operational problems. Chapter 4 presents the overall design of the system, while Chapter 5 explains the database structure together with the semantic enhancement design. Chapter 6 further describes the detailed design of key parts such as permission control, inventory processing, workflow linkage and semantic support. Chapter 7 shows the implementation results, Chapter 8 reports the testing work, and Chapter 9 concludes the thesis and discusses the current limitations together with possible future improvements.",
    725: "本文不是从一般性的数字化概念出发去讨论系统建设，而是把关注点放在南派食品这类中小制造企业的实际问题上。围绕信息分散、库存台账不清、批次追溯困难、权限边界模糊以及系统后续维护成本偏高等问题，论文逐步完成了需求分析、总体设计、数据库设计、详细设计、系统实现和阶段性测试，并将结果落实到可运行的系统原型之中。",
    726: "从当前完成情况看，系统已经形成了较明确的工程骨架。前端由基座、人事、物料、应用中心和移动端几个核心模块协同运行，后端则以数据库中心架构承载主要数据结构、部分业务约束和权限边界，流程配置与审批联动也已经完成阶段性落地。数据库、接口与运行时服务通过 Docker Compose 统一组织，多个前端模块通过 PM2 管理运行，这使系统在部署和维护方式上都比较清晰。",
    727: "在架构选择上，本文没有继续沿用传统的“前端 + 独立应用服务 + 数据库”三层做法，而是把核心数据、权限控制和部分业务逻辑尽量下沉到数据库侧，再借助 PostgREST 压缩常规接口开发量。对中小制造企业来说，这样的取舍更贴合“部署负担不能太重、后续修改不能太麻烦”的现实条件，也构成了本文系统实现的重要特点。",
    730: "论文的另一项特点，是把轻量化语义机制真正放进了系统实现过程，而不是停留在概念说明层面。通过表级语义、列级语义、语义中文化回填和本体关系工作台，系统开始能够对业务表、字段和表间关系做更统一的解释。它的价值不在于追求复杂推理，而在于为流程、权限和动态配置提供一致的表达基础。",
    732: "整体来看，本文已经完成了从需求分析到实现验证的主要链条，也证明了数据库中心架构、微前端拆分、流程支撑和轻量化语义机制在中小制造企业场景下具备可行性。以南派食品为例，这套系统虽然还没有覆盖企业全部业务方向，但已经把最核心的几条链路稳定下来，具备继续扩展的基础。",
    758: "本次毕业设计从选题、调研到系统实现和论文整理，前后经历了较长过程。在这一过程中，我得到了老师、同学、企业相关人员以及家人的帮助。论文完成之际，我想对这些给予我支持的人表示感谢。",
    759: "首先要感谢我的指导教师郑苑丹老师。无论是在课题方向的确定、论文结构的调整，还是系统实现过程中的问题分析与修改上，老师都给了我很多具体而耐心的指导。许多看似细小的意见，实际上都直接影响了论文最后的完整性和表达质量。",
    760: "同时，也要感谢在项目开发和论文整理过程中给予我帮助的同学和朋友。在联调、资料收集、问题讨论和文稿修改阶段，大家提供了很多实际支持，这些帮助让我在推进过程中少走了不少弯路。",
    761: "我还要感谢广东南派食品有限公司在业务调研阶段提供的场景资料与业务信息。正因为有真实企业流程作为背景，这次毕业设计才没有停留在纯概念设计层面，而能够围绕中小制造企业信息化建设中的真实问题展开分析和实现。",
    762: "此外，也感谢大学四年中所有给予我帮助和启发的任课老师。软件工程、数据库系统、程序设计、前端开发以及系统分析与设计等课程，为本次毕业设计提供了持续的知识积累和方法基础。",
    763: "最后，感谢一直支持我的家人。正是他们在整个毕业设计期间给予的理解和陪伴，让我能够比较稳定地完成系统开发和论文写作。",
    764: "毕业设计的结束并不意味着学习的结束。通过这次课题实践，我更清楚地认识到把课堂知识落到真实场景并不容易，但也正是在这个过程中，系统分析、设计、实现和写作能力才真正被串联起来。"
}


def main() -> None:
    doc = Document(str(SRC))
    changed = 0
    for idx, new_text in PARA_REPLACEMENTS.items():
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].text = new_text
            changed += 1
    doc.save(str(OUT))
    doc.save(str(OUT_EN))
    print(f"saved: {OUT}")
    print(f"saved: {OUT_EN}")
    print(f"changed paragraphs: {changed}")


if __name__ == "__main__":
    main()
