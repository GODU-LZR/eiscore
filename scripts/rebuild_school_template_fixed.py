from pathlib import Path
import re
import shutil
import tempfile
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

TEMPLATE = Path(r"C:/Users/Twist/Desktop/school_template_2015_design.docx")
MD = next(Path('docs').glob('*总稿_目录版*.md'))
OUT = Path('docs') / '毕业论文总稿_按校模板终稿_2026-03-15.docx'

TITLE_CN = '面向湛江中小制造企业的企业信息化系统核心设计与实现——以广东南派食品有限公司为例'
TITLE_EN = 'Core Design and Implementation of an Enterprise Informatization System for Small and Medium-Sized Manufacturing Enterprises in Zhanjiang: A Case Study of Guangdong Nanpai Food Co., Ltd.'
AUTHOR = '林志荣'
STUDENT_ID = '202111701318'
COLLEGE = '数学与计算机学院'
CLASS_NAME = '软件1223'
MAJOR = '软件工程'
DEGREE = '工学学士'
ADVISOR = '郑苑丹'
ADVISOR_TITLE = '讲师'
SUB_ADVISOR = ''
SUB_ADVISOR_TITLE = ''
DEFENSE_DATE = '2026年     月     日'
KEYWORDS_CN = '中小制造企业；企业信息化；微前端；数据库中心架构；轻量本体语义增强'
KEYWORDS_EN = 'small and medium-sized manufacturing enterprises; enterprise informatization; micro-frontends; database-centric architecture; lightweight ontology semantic enhancement'

text = MD.read_text(encoding='utf-8')


def section_between(start, end):
    s = text.index(start) + len(start)
    e = text.index(end, s)
    return text[s:e].strip('\n')


cn_summary = section_between('# 设计总说明', '# Design Summary')
en_summary = section_between('# Design Summary', '# 第一章 绪论')
body = text[text.index('# 第一章 绪论'):text.index('# 致谢')].strip('\n')
ack = section_between('# 致谢', '# 参考文献')
refs = text[text.index('# 参考文献') + len('# 参考文献'):].strip('\n')


def parse_blocks(section_text):
    ls = section_text.splitlines()
    blocks = []
    i = 0
    while i < len(ls):
        s = ls[i].strip()
        if not s:
            i += 1
            continue
        if s.startswith('|'):
            rows = []
            while i < len(ls) and ls[i].strip().startswith('|'):
                rows.append(ls[i].strip())
                i += 1
            parsed = []
            for ridx, row in enumerate(rows):
                if ridx == 1 and set(row.replace('|', '').replace('-', '').replace(':', '').strip()) == set():
                    continue
                parsed.append([c.strip() for c in row.strip('|').split('|')])
            blocks.append(('table', parsed))
            continue
        m = re.match(r'^(#{1,4})\s+(.*)$', s)
        if m:
            blocks.append(('heading', len(m.group(1)), m.group(2).strip().replace('`', '')))
            i += 1
            continue
        m = re.match(r'^(\d+)\.\s+(.*)$', s)
        if m:
            blocks.append(('list', f"{m.group(1)}. {m.group(2)}"))
            i += 1
            continue
        blocks.append(('para', s.replace('`', '')))
        i += 1
    return blocks


cn_blocks = parse_blocks(cn_summary) + [('para', f'关键词：{KEYWORDS_CN}')]
en_blocks = parse_blocks(en_summary) + [('para', f'Keywords: {KEYWORDS_EN}')]
body_blocks = parse_blocks(body)
ack_blocks = parse_blocks(ack)
ref_blocks = [line.strip() for line in refs.splitlines() if line.strip()]

doc = Document(str(TEMPLATE))


def set_run_fonts(run, size=10.5, bold=None):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def find_para(txt, occurrence=1):
    count = 0
    for p in doc.paragraphs:
        if p.text.strip() == txt:
            count += 1
            if count == occurrence:
                return p
    raise ValueError(txt)


def remove_paragraph_numbering(para):
    p_pr = para._p.get_or_add_pPr()
    num_pr = p_pr.find(qn('w:numPr'))
    if num_pr is not None:
        p_pr.remove(num_pr)


if len(doc.tables) >= 3:
    doc.tables[0].cell(0, 0).text = TITLE_CN
    doc.tables[1].cell(0, 0).text = TITLE_EN
    t = doc.tables[2]
    t.cell(0, 1).text = AUTHOR
    t.cell(0, 4).text = STUDENT_ID
    t.cell(1, 1).text = COLLEGE
    t.cell(1, 5).text = CLASS_NAME
    t.cell(2, 1).text = MAJOR
    t.cell(3, 1).text = DEGREE
    t.cell(4, 1).text = ADVISOR
    t.cell(4, 5).text = ADVISOR_TITLE
    t.cell(5, 1).text = SUB_ADVISOR
    t.cell(5, 5).text = SUB_ADVISOR_TITLE
    t.cell(6, 1).text = DEFENSE_DATE

for p in list(doc.paragraphs):
    if '请按《广东海洋大学本科毕业论文' in p.text:
        p._element.getparent().remove(p._element)
        break

find_para('设计总说明').text = '设计总说明'
find_para('introduction').text = 'Introduction'
find_para('（毕业设计题目，黑体三号，加粗，行距为固定值20磅，段前段后各1行）').text = TITLE_CN
find_para('（专业，学号，姓名，中间以逗号分隔，宋体五号，行距为固定值20磅）').text = f'{MAJOR}，{STUDENT_ID}，{AUTHOR}'
find_para('指导教师：（教师姓名，宋体五号，行距为固定值20磅）').text = f'指导教师：{ADVISOR}'

for p in list(doc.paragraphs):
    if p.text.strip() == '附  录':
        elem = p._element
        while elem is not None:
            nxt = elem.getnext()
            elem.getparent().remove(elem)
            elem = nxt
        break


def delete_between(start_para, end_para):
    elem = start_para._element.getnext()
    while elem is not None and elem is not end_para._element:
        nxt = elem.getnext()
        elem.getparent().remove(elem)
        elem = nxt


def insert_paragraph_after(element, text='', style='Normal', no_indent=False):
    new_p = OxmlElement('w:p')
    element.addnext(new_p)
    para = Paragraph(new_p, doc._body)
    para.style = style
    if no_indent:
        para.paragraph_format.first_line_indent = Pt(0)
    for part in re.split(r'(`[^`]+`)', text):
        if not part:
            continue
        if part.startswith('`') and part.endswith('`'):
            part = part[1:-1]
        run = para.add_run(part)
        set_run_fonts(run)
    return para


def insert_table_after(element, rows):
    table = doc.add_table(rows=0, cols=max(len(r) for r in rows))
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for cidx in range(len(cells)):
            cells[cidx].text = row[cidx] if cidx < len(row) else ''
            for p in cells[cidx].paragraphs:
                for r in p.runs:
                    set_run_fonts(r, bold=(ridx == 0))
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    element.addnext(tbl)
    return tbl


def insert_blocks_before(marker_para, blocks):
    prev = marker_para._element.getprevious()
    if prev is None:
        prev = marker_para._element
    for block in blocks:
        kind = block[0]
        if kind == 'heading':
            _, level, txt = block
            style = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4'}.get(level, 'Normal')
            p = insert_paragraph_after(prev, txt, style=style, no_indent=True)
            remove_paragraph_numbering(p)
            prev = p._element
        elif kind == 'para':
            p = insert_paragraph_after(prev, block[1], style='Normal')
            prev = p._element
        elif kind == 'list':
            p = insert_paragraph_after(prev, block[1], style='Normal', no_indent=True)
            prev = p._element
        elif kind == 'table':
            prev = insert_table_after(prev, block[1])
            p = insert_paragraph_after(prev, '', style='Normal', no_indent=True)
            prev = p._element


p_cn = find_para('设计总说明')
p_en = find_para('Introduction')
delete_between(p_cn, p_en)
insert_blocks_before(p_en, cn_blocks)

p_title = find_para(TITLE_CN)
delete_between(p_en, p_title)
insert_blocks_before(p_title, en_blocks)

p_body_anchor = find_para('毕业设计说明书')
p_ack = find_para('鸣  谢')
delete_between(p_body_anchor, p_ack)
insert_blocks_before(p_ack, body_blocks)

p_ref = find_para('参考文献')
delete_between(p_ack, p_ref)
insert_blocks_before(p_ref, ack_blocks)

for sib in list(p_ref._element.itersiblings()):
    sib.getparent().remove(sib)
prev = p_ref._element
for ref in ref_blocks:
    p = insert_paragraph_after(prev, ref, style='Normal', no_indent=True)
    prev = p._element

level_map = {'Heading 1': 0, 'Heading 2': 1, 'Heading 3': 2, 'Heading 4': 3}

for p in doc.paragraphs:
    style_name = p.style.name
    if style_name in level_map:
        p.style = doc.styles['Normal']
        p.paragraph_format.first_line_indent = 0
        remove_paragraph_numbering(p)
        p_pr = p._p.get_or_add_pPr()
        outline = p_pr.find(qn('w:outlineLvl'))
        if outline is None:
            outline = OxmlElement('w:outlineLvl')
            p_pr.append(outline)
        outline.set(qn('w:val'), str(level_map[style_name]))
    elif p.text.strip() in {'?????', 'Introduction'}:
        remove_paragraph_numbering(p)
    for r in p.runs:
        if style_name in level_map:
            set_run_fonts(r, size=12, bold=True)
        else:
            set_run_fonts(r)

doc.save(str(OUT))


def update_toc_via_word(docx_path):
    try:
        import win32com.client  # type: ignore
    except Exception as exc:
        print(f'WARN: win32com unavailable, skip TOC update: {exc}')
        return

    temp_docx = Path(tempfile.gettempdir()) / docx_path.name
    shutil.copy2(docx_path, temp_docx)

    word = None
    doc_com = None
    try:
        word = win32com.client.DispatchEx('Word.Application')
        word.Visible = False
        word.DisplayAlerts = 0
        doc_com = word.Documents.Open(str(temp_docx.resolve()))
        doc_com.Repaginate()
        for toc in doc_com.TablesOfContents:
            toc.Update()
        doc_com.Fields.Update()
        doc_com.Save()
        doc_com.Close(False)
        doc_com = None
        word.Quit()
        word = None
        shutil.copy2(temp_docx, docx_path)
    finally:
        if doc_com is not None:
            doc_com.Close(False)
        if word is not None:
            word.Quit()


update_toc_via_word(OUT)
print(OUT)
