from __future__ import annotations

import re
from pathlib import Path

from docx import Document


DESKTOP_PAPER = Path(r"C:\Users\Twist\Desktop\论文")
SOURCE = DESKTOP_PAPER / "主稿" / "毕业论文初稿v2.8_降风险修订版_2026-04-06.docx"
TARGET = DESKTOP_PAPER / "主稿" / "毕业论文初稿v2.8_降风险终修版_2026-04-06.docx"


def clean_caption_suffixes(text: str) -> str:
    text = re.sub(r"(图\d+-\d+[^\n。；]*?)\.(svg|png)\b", r"\1", text, flags=re.IGNORECASE)
    text = re.sub(r"(表\d+-\d+[^\n。；]*?)\.(svg|png)\b", r"\1", text, flags=re.IGNORECASE)
    return text


def rewrite(text: str) -> str:
    t = clean_caption_suffixes(text)

    replacements = [
        (
            "图中从企业现实约束出发，先落到需求分析，再进入总体设计、关键技术路线、系统实现和测试验证，最后回到说明书成文。这样安排的目的，是把“南派食品到底遇到了什么问题”“EISCore 为什么这样设计”“后面各章分别承接哪一段工作”放在同一张图里讲清楚。见图1-1。",
            "图1-1主要交代本设计的展开顺序。阅读时可以先看到企业现实约束，再看到需求分析、系统设计、实现与验证几个阶段如何逐步展开。这样处理的目的很直接：先把问题来源交代清楚，再说明后续各章分别承接哪一段工作。见图1-1。",
        ),
        (
            "本设计组织结构",
            "设计说明书结构安排",
        ),
        (
            "本设计共分为八章，各章内容安排如下。第一章为绪论，主要介绍课题研究背景、研究意义、国内外研究现状以及本设计的主要研究内容与结构安排。第二章介绍系统开发涉及的关键技术和工具，包括 Vue 3、微前端、PostgreSQL、PostgREST、BPMN 以及本体语义相关技术。第三章进行系统需求分析，围绕用户角色、核心业务流程和非功能需求展开说明。第四章给出系统概要设计，重点说明系统架构、能力组织和关键设计思路。第五章介绍数据库设计与语义增强设计，包括数据模型、主要表结构和本体语义组织方式。第六章对核心实现链路进行详细设计，分别说明人事管理、物料管理、流程联动和动态应用配置等部分的实现机制。第七章展示系统实现效果，并结合主要页面和业务链路说明系统运行情况。第八章对系统进行测试与分析，并在结论部分总结全文工作、指出存在的不足以及未来改进方向。",
            "本设计后续内容按“需求、概要、数据、详细、实现与验证”这一顺序展开。前文先把背景、现状和技术基础交代清楚；随后进入系统需求分析、系统概要设计和数据库设计；再在详细设计中说明关键实现链路；最后回到系统实现、测试结果和当前边界。这样安排比单纯罗列章节更便于对应每一部分的作用。",
        ),
        (
            "因此，本节关注的不是抽象本体如何继续扩展，而是语义能力怎样嵌入现有系统。对当前项目来说，轻量化语义主要承担三件事：给页面补充对象解释、给应用中心补充动态表结构说明、给受控工具提供对象类型判断依据。",
            "因此，本节不讨论抽象本体还能怎样继续扩展，而只说明语义能力在当前系统里的落点。就现阶段实现而言，它主要做三件事：一是补充页面对象解释，二是补充应用中心里的动态表结构说明，三是给受控工具提供对象类型判断依据。",
        ),
        (
            "图6-3展示了流程启动和状态写回的处理链路。当前端从业务页面或应用中心发起流程时，请求会先经由 PostgREST 进入数据库函数；`workflow.start_workflow_instance` 负责创建实例并写入初始状态，后续节点推进则由 `workflow.transition_workflow_instance` 接手。推进前，系统会先结合 `workflow.can_execute_task` 和任务分派规则校验当前操作者；校验通过后，再更新流程实例、记录审批日志，并依据 `app_center.workflow_state_mappings` 把结果写回对应业务表。见图6-3。",
            "图6-3展示了流程启动和状态写回的处理链路。当前端从业务页面或应用中心发起流程时，请求先经过 PostgREST，再进入数据库函数。`workflow.start_workflow_instance` 用于生成流程实例并写入初始状态；后续节点推进由 `workflow.transition_workflow_instance` 处理。推进之前，系统会先结合 `workflow.can_execute_task` 与任务分派规则判断当前操作者是否具备继续执行资格。只有校验通过，才会更新流程实例、写入审批日志，并依据 `app_center.workflow_state_mappings` 把状态结果回写到对应业务表。见图6-3。",
        ),
        (
            "在完成采购入库与来料检验用例规约之后，本设计进一步绘制对应活动图。采购入库与来料检验活动图如图3-7所示。该图重点反映“到货登记—来料检验—是否合格—正式入库/待补录”这条链路，对应采购与质检之间的协同需求，也与表3-4中的用例规约形成一一对应。见图3-7。",
            "表3-4先把采购入库与来料检验场景中的参与角色、异常分支和后置条件固定下来，随后再用活动图说明其流程走向。图3-7对应的就是这条链路，重点表现到货登记、来料检验、合格判定和正式入库之间的先后关系。见图3-7。",
        ),
        (
            "在完成生产领料与补料用例规约之后，本设计进一步绘制对应活动图。生产领料与补料活动图如图3-8所示。该图重点体现标准领料与补料场景的差别，以及补料业务为什么必须作为独立需求存在。对于食品加工企业而言，原料属性波动会直接影响配方执行，因此“领料完成后仍可能继续补料”是一个具有行业特征的业务分支。见图3-8。",
            "表3-5先说明生产领料与补料场景中的参与角色、触发条件和异常处理，再由图3-8补充其流程走向。该图重点表现标准领料与补料分支为什么要分开处理，也说明食品加工场景中“领料结束后仍可能继续补料”的业务特点。见图3-8。",
        ),
        (
            "此外，项目在开发过程中还使用了基于浏览器的 code-server 作为辅助开发环境。它不属于系统交付后的核心运行链路，但在多人协同调试、远程接入和临时修改场景中确实发挥了作用，因此在本节一并说明。",
            "此外，项目开发过程中还用到了基于浏览器的 code-server。它不属于系统交付后的核心运行链路，但在多人协同调试、远程接入和临时修改场景中确实发挥过作用，因此在本节顺带交代。",
        ),
    ]

    for old, new in replacements:
        if old in t:
            t = t.replace(old, new)

    if "图1-1 论文研究思路与技术路线图" in t and "设计说明书结构安排" not in t:
        t = t.replace("图1-1 论文研究思路与技术路线图", "图1-1 论文研究思路与技术路线图\n设计说明书结构安排")

    # tighten a few common high-risk phrasing patterns
    t = t.replace("这样安排的目的，是", "这样安排主要是")
    t = t.replace("重点表现", "主要表现")
    t = t.replace("其流程走向", "流程走向")

    return t


def process_document(src: Path, dst: Path) -> None:
    doc = Document(str(src))
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            new_text = rewrite(para.text)
            if new_text != para.text:
                para.text = new_text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text:
                        new_text = clean_caption_suffixes(para.text)
                        if new_text != para.text:
                            para.text = new_text

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))


if __name__ == "__main__":
    process_document(SOURCE, TARGET)
    print(TARGET)
