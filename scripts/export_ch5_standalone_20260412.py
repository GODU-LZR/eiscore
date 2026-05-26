from __future__ import annotations

import importlib.util
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


REPO_ROOT = Path(__file__).resolve().parents[1]
DESKTOP = Path.home() / "Desktop"
THESIS_ROOT = next(
    p for p in DESKTOP.iterdir() if p.is_dir() and "论文" in p.name
)
OUT_DOC = THESIS_ROOT / "主稿" / "第五章_数据库设计_单独版_2026-04-12.docx"


def load_ch5_module():
    script_path = REPO_ROOT / "scripts" / "rewrite_ch5_db_design_20260412.py"
    spec = importlib.util.spec_from_file_location("rewrite_ch5", script_path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def format_run(run, size=12, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def style_paragraph(paragraph, size=12, bold=False, center=False):
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        format_run(run, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    if level == 1:
        p.style = "Heading 1"
        size = 16
    elif level == 2:
        p.style = "Heading 2"
        size = 14
    else:
        p.style = "Heading 3"
        size = 12
    p.add_run(text)
    style_paragraph(p, size=size, bold=True)
    return p


def add_body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    try:
        p.style = "Body Text"
    except Exception:
        pass
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    style_paragraph(p, size=12)
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)
    style_paragraph(p, size=10.5)
    return p


def set_cell_text(cell, text: str):
    cell.text = str(text)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, size=10.5)


def add_table(doc: Document, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            set_cell_text(table.cell(i, j), value)
    return table


def add_table_section(doc: Document, title: str, intro: str, caption: str, explain: str, rows):
    add_heading(doc, title, 3)
    add_body(doc, intro)
    add_caption(doc, caption)
    add_body(doc, explain)
    add_table(doc, rows)


def build_doc() -> Document:
    mod = load_ch5_module()
    src = Document(str(mod.OUT_DOC)) if Path(mod.OUT_DOC).exists() else None
    doc = Document()
    sec = doc.sections[0]
    if src:
        src_sec = src.sections[0]
        sec.top_margin = src_sec.top_margin
        sec.bottom_margin = src_sec.bottom_margin
        sec.left_margin = src_sec.left_margin
        sec.right_margin = src_sec.right_margin
        sec.page_width = src_sec.page_width
        sec.page_height = src_sec.page_height
    else:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.17)
        sec.right_margin = Cm(3.17)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("5 数据库设计")
    style_paragraph(p, size=16, bold=True)

    add_heading(doc, "5.1 数据库设计目标", 2)
    add_body(doc, "数据库设计是本设计的核心部分。与传统企业系统中数据库主要承担持久化存储不同，EISCore 采用数据库中心架构，因此数据库不仅要完成数据保存，还要承担权限控制、流程关联、状态映射和语义增强等职责。换句话说，数据库设计既是系统的数据模型设计，也是系统治理能力设计的一部分。")
    add_body(doc, "结合南派食品的业务需求和当前项目实现情况，本系统数据库设计主要围绕以下目标展开。第一，建立稳定的基础主数据模型，统一组织架构、用户、角色、物料、仓库等关键对象。第二，围绕库存批次、库存流水、入库出库和盘点等业务构建仓储链路，支撑批次追溯和库存透明化。第三，为应用中心和工作流运行建立独立的数据域，使动态应用、流程设计和审批能力能够在同一数据库中得到统一管理。第四，通过 RLS 和语义增强结构，为后续权限治理和本体语义描述提供底层支撑。")

    add_heading(doc, "5.2 数据库总体结构设计", 2)
    add_body(doc, "EISCore 数据库以 PostgreSQL 16 为核心，采用多 Schema 方式对不同业务域进行划分。当前系统中与本设计主线密切相关的 Schema 主要包括 public、hr、scm、app_center、workflow 和 app_data。")
    add_body(doc, "其中，public Schema 主要保存组织结构、用户、角色、权限、物料主数据及部分系统公共表；hr Schema 主要承载员工档案、考勤、工资和人事扩展信息；scm Schema 主要承载仓库、库存批次、库存流水和盘点等供应链与仓储数据；app_center Schema 主要承载应用中心中的应用注册、发布路由、状态映射和执行日志；workflow Schema 主要承载流程定义、流程实例和任务分派；app_data Schema 则用于承载动态生成的数据应用表。")
    add_body(doc, "这种划分方式的优势主要有三点。第一，不同业务域在逻辑上更加清晰，便于后续维护和权限分层。第二，流程、应用中心和业务主数据之间虽然相互关联，但又能保持清楚的边界，避免所有表混杂在同一命名空间下。第三，配合 PostgREST 和 RLS 后，不同 Schema 能够更方便地形成权限隔离和接口边界。")

    add_heading(doc, "5.3 核心数据实体关系设计", 2)
    add_body(doc, "从本设计主线出发，EISCore 的核心实体关系主要围绕组织权限主线、物料仓储主线和应用流程主线展开，其核心实体关系如图5-1所示。")
    fig_er = Path(mod.FIG_ER)
    if fig_er.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(fig_er), width=Inches(5.8))
    add_caption(doc, "图5-1 EISCore核心数据实体关系图")
    add_body(doc, "在组织权限主线中，departments、positions、users、roles 和 user_roles 构成了基础的组织授权关系。部门表支持层级结构，岗位表和用户表分别通过外键关联部门，角色表承载系统权限分类，用户与角色之间再通过关联表形成多对多关系。这条主线决定了用户能看到什么菜单、能处理哪些任务、能访问哪些数据域。")
    add_body(doc, "在物料与仓储主线中，raw_materials、warehouses、inventory_batches 和 inventory_transactions 构成了库存业务链路。物料主数据表定义所有基础物料对象，仓库表通过树形结构描述仓库、库区和库位，库存批次表记录某个物料在某个位置上的批次库存，库存流水表则记录所有库存变化行为。通过这条主线，系统可以支撑采购入库、生产领料、补料出库、生产入库、销售出库以及盘点调整。")
    add_body(doc, "在应用与流程主线中，categories、apps、workflow.definitions、workflow.instances、workflow.task_assignments、published_routes 和 workflow_state_mappings 构成了动态应用与流程协同能力的核心结构。应用表用于登记动态应用和流程应用，流程定义表与应用表建立关联，流程实例表和任务分派表承载运行态，而发布路由和状态映射表则分别用于前端挂载与业务状态写回。三条主线共同构成 EISCore 的数据库骨架。")

    add_heading(doc, "5.4 组织与权限相关表设计", 2)
    add_heading(doc, "5.4.1 部门表与岗位表", 3)
    add_body(doc, "部门表 public.departments 用于描述组织结构，主要字段包括 id、name、parent_id、leader_id、sort 和状态相关字段。通过 parent_id 自关联，系统可以表达树形组织结构，为人事归属和数据范围控制提供依据。岗位表 public.positions 则通过 dept_id 关联部门，描述部门内的岗位信息。")
    add_body(doc, "这种设计的意义不只在于展示组织架构。对于南派食品当前场景，仓库人员、质检人员、销售文员和管理人员分属不同部门和岗位，部门与岗位信息会进一步参与角色配置、流程候选人筛选以及数据范围约束，因此组织表是权限治理的起点。")
    add_heading(doc, "5.4.2 用户表、角色表与用户角色关联表", 3)
    add_body(doc, "用户表 public.users 用于保存系统登录用户及其基本信息，关键字段包括 username、full_name、dept_id 和 position_id。角色表 public.roles 保存角色编码、角色名称以及所属部门等信息。由于同一用户在系统里可能承担审批、仓储、查看经营数据等多种职责，因此系统通过 public.user_roles 建立用户与角色之间的多对多关系。")
    add_body(doc, "这一设计方式既能支持“一个用户一个主角色”的简单场景，也能支持“同一用户在不同业务中承担不同角色”的复杂场景。它与第六章中权限控制链路的实现保持一致：前端完成身份识别后，PostgREST 和 PostgreSQL 进一步根据角色声明和 RLS 策略裁剪数据访问范围。")

    add_heading(doc, "5.5 物料与仓储相关表设计", 2)
    add_heading(doc, "5.5.1 物料主数据表", 3)
    add_body(doc, "物料主数据是仓储链路的起点。当前系统中的 public.raw_materials 虽然命名为原料表，但在实现上承担了更广义的物料基础对象存储职责。除了名称、分类、部门归属等字段外，该表还通过 properties 字段兼容非标准扩展属性，用于适应食品加工场景中“同名物料、属性波动”的情况。")
    add_heading(doc, "5.5.2 仓库表与层级结构", 3)
    add_body(doc, "仓库表 scm.warehouses 采用树形结构设计，通过 parent_id 同时表示仓库、库区和库位三个层级。这一做法比把仓库、区域、库位拆成三张表更适合当前项目的规模，也便于移动端盘点和冷库模式复用同一套位置数据。对于食品企业常见的“一个仓库内再细分多个区域和位置”的管理方式，这种设计更贴近现场。")
    add_heading(doc, "5.5.3 库存批次表", 3)
    add_body(doc, "库存批次表 scm.inventory_batches 是当前数据库设计中最关键的业务表之一。该表以 material_id 和 warehouse_id 为基础，记录批次号、可用数量、锁定数量、生产日期、过期日期以及批次状态等信息。南派食品存在保质期管理、批次追溯和检验延迟等业务约束，因此这张表承担了库存台账和追溯链条中的关键角色。")
    add_heading(doc, "5.5.4 库存流水表", 3)
    add_body(doc, "库存流水表 scm.inventory_transactions 记录所有库存变化行为，包括采购入库、生产领料、补料出库、生产入库、销售出库和盘点调整等。它通过 material_id、batch_id 和 warehouse_id 与物料、批次和仓库建立关联，并保留交易类型、数量、关联单据和操作时间等信息。该表既是库存审计的依据，也是后续批次追溯、问题回放和状态对账的基础。")

    add_heading(doc, "5.6 应用中心与流程相关表设计", 2)
    add_heading(doc, "5.6.1 应用分类表与应用表", 3)
    add_body(doc, "应用分类表 app_center.categories 和应用表 app_center.apps 共同构成应用中心的基础结构。分类表负责区分数据应用、流程应用和其他应用类型；应用表保存应用名称、分类、类型、状态、配置以及 BPMN XML 等信息。由于 EISCore 并不只是固定页面集合，还需要支持动态应用配置、闪念应用草稿和流程设计，因此 app_center.apps 在数据库中扮演了“应用注册中心”的角色。")
    add_heading(doc, "5.6.2 发布路由与状态映射表", 3)
    add_body(doc, "发布路由表 app_center.published_routes 用于记录应用发布后的访问路径，便于前端基座动态挂载；状态映射表 app_center.workflow_state_mappings 用于把流程任务节点与具体业务表的状态字段建立关联，使流程推进时能够驱动业务状态写回。基座挂载页面时会读取前者，流程推进和写回业务状态时会读取后者，因此这两张表都直接参与系统运行。")
    add_heading(doc, "5.6.3 流程定义表、实例表与任务分派表", 3)
    add_body(doc, "流程定义表 workflow.definitions 用于保存 BPMN 流程定义，并通过 app_id 关联应用中心中的流程应用；流程实例表 workflow.instances 保存流程运行中的实例状态，包括当前任务节点、业务键和状态等；任务分派表 workflow.task_assignments 用于定义某个流程节点可由哪些角色或用户处理。通过这三类表，系统能够从静态流程定义延伸到动态流程运行与任务权限分派。")

    add_heading(doc, "5.7 数据安全与权限控制设计", 2)
    add_body(doc, "本设计的数据安全控制并不依赖厚重的后端服务层，而是通过 PostgREST、JWT 声明和 PostgreSQL 的 RLS 策略共同落地。应用请求进入数据库前，PostgREST 会携带身份声明；数据库中的 RLS 策略再根据用户角色、部门范围和表级规则决定是否允许读取、写入或推进流程。对当前这种多子应用共享同一数据库的实现方式而言，这种“接口声明 + 数据库裁剪”的组合更容易保证边界一致。")
    add_body(doc, "例如，应用中心中的 apps、published_routes、workflow_state_mappings 和 execution_logs 等表均配置了针对不同操作类型的访问限制；工作流相关表也通过角色声明约束写入和管理权限。这样做的直接好处，是让菜单可见性、按钮权限和数据写入边界不再各自为政，而是回到同一套数据库治理规则中。")

    add_heading(doc, "5.8 语义增强相关设计", 2)
    add_body(doc, "轻量化语义增强并不是在数据库之外再搭一个独立知识库，而是在现有数据结构上增加一层可解释描述能力。当前设计主要围绕表级语义、列级语义和关系说明展开，用于支撑本体关系工作台、闪念应用对象解释以及数据应用字段理解。它的目标不是替代业务表，而是让应用中心、数据表格应用和受控 Agent 运行时能够更准确地理解“这张表是什么、字段代表什么、关系怎样组合”。")
    add_body(doc, "从数据库设计角度看，语义增强的重点在于保留稳定的元数据锚点，而不是引入复杂推理引擎。也就是说，语义层必须依附现有的组织、物料、流程和应用表而存在，只有这样，后续在动态应用配置和对象解释中才能保持一致。")

    add_heading(doc, "5.9 关键物理表结构设计", 2)
    add_body(doc, "按照毕业设计说明书对数据库设计章节的要求，在说明总体结构和实体关系之后，还需要进一步给出关键物理表的字段结构。结合当前项目实现情况，本设计优先选取用户、角色、物料、库存、应用中心和流程运行等对主线影响最大的表进行展示。各表在本章中均先给出正文引入，再列出字段结构，避免只堆表格而缺少解释。")
    add_table_section(doc, "5.9.1 用户表", "为说明系统核心组织权限数据的基础结构，用户表的主要字段设计如表5-1所示。", "表5-1 用户表", "表5-1 展示了 public.users 的主要字段结构。该表承担账号信息、基础身份属性和组织归属信息的存储功能，是系统登录、身份识别和后续 RLS 数据裁剪的重要入口。", mod.TABLE_5_1)
    add_table_section(doc, "5.9.2 角色表", "为说明系统角色与权限分类的基础数据结构，角色表的主要字段设计如表5-2所示。", "表5-2 角色表", "表5-2 展示了 public.roles 的主要结构。该表用于定义系统角色编码、角色名称以及角色说明信息，并可与部门形成关联，为后续权限分配和流程节点候选角色配置提供基础。", mod.TABLE_5_2)
    add_table_section(doc, "5.9.3 用户角色关联表", "为说明用户与角色之间的关联方式，用户角色关联表的主要字段设计如表5-3所示。", "表5-3 用户角色关联表", "表5-3 展示了 public.user_roles 的字段结构。该表用于建立用户与角色之间的多对多关系，使系统能够支持一个用户对应多个角色，满足审批、仓储和管理查看等复合授权场景。", mod.TABLE_5_3)
    add_table_section(doc, "5.9.4 物料主数据表", "为说明物料主数据在系统中的组织方式，物料主数据表的主要字段设计如表5-4所示。", "表5-4 物料主数据表", "表5-4 展示了 public.raw_materials 的主要字段。该表虽然命名为原料表，但在当前系统实现中承担了较广义的物料基础信息存储职责，并通过 properties 字段兼容扩展属性。", mod.TABLE_5_4)
    add_table_section(doc, "5.9.5 仓库表", "为说明仓库基础信息的存储方式，仓库表的主要字段设计如表5-5所示。", "表5-5 仓库表", "表5-5 展示了 scm.warehouses 的主要结构。该表采用树形结构设计，可同时表示仓库、库区和库位三个层级，适合食品企业中分区管理和现场布局扩展的实际需要。", mod.TABLE_5_5)
    add_table_section(doc, "5.9.6 库存批次表", "为说明库存批次管理与保质期追踪的数据基础，库存批次表的主要字段设计如表5-6所示。", "表5-6 库存批次表", "表5-6 展示了 scm.inventory_batches 的字段结构。该表是本系统中最重要的仓储业务表之一，主要用于保存某种物料在某个仓位上的批次库存状态，为批次追溯、保质期管理和台账统计提供支撑。", mod.TABLE_5_6)
    add_table_section(doc, "5.9.7 库存流水表", "为说明库存收发流水的记录结构，库存流水表的主要字段设计如表5-7所示。", "表5-7 库存流水表", "表5-7 展示了 scm.inventory_transactions 的字段结构。该表记录所有出入库和库存变化行为，是连接库存状态与业务动作的核心表。系统通过该表保存单据号、业务类型、数量变化和关联单据信息，为库存审计和问题追溯提供依据。", mod.TABLE_5_7)
    add_table_section(doc, "5.9.8 应用注册表", "为说明应用中心中已注册应用的组织方式，应用注册表的主要字段设计如表5-8所示。", "表5-8 应用注册表", "表5-8 展示了 app_center.apps 的字段结构。该表用于保存应用中心中的应用定义信息，同时支持数据应用、流程应用、Flash 草稿应用和自定义应用，是动态应用管理能力的核心载体。", mod.TABLE_5_8)
    add_table_section(doc, "5.9.9 流程实例表", "为说明流程运行过程中的实例数据结构，流程实例表的主要字段设计如表5-9所示。", "表5-9 流程实例表", "表5-9 展示了 workflow.instances 的字段结构。该表用于保存流程运行态数据，是系统流程引擎与业务状态联动的核心表之一。与只保存设计稿的流程定义表不同，该表更关注具体业务对象当前运行到哪个任务节点以及实例状态如何变化。", mod.TABLE_5_9)

    add_heading(doc, "5.10 本章小结", 2)
    add_body(doc, "本章围绕 EISCore 的数据库设计展开，依次说明了数据库设计目标、总体结构、核心实体关系、组织权限主线、物料仓储主线以及应用流程主线，并补充给出了九张关键物理表结构。整体来看，本系统数据库不仅承担数据存储职能，还承担权限控制、流程协同和语义增强等治理能力。这种数据库中心设计方式与本设计“轻量、可扩展、低运维成本”的总体目标保持一致，也为后续系统详细设计提供了稳定的数据基础。")
    return doc


def main():
    doc = build_doc()
    OUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOC))
    print(OUT_DOC)


if __name__ == "__main__":
    main()
