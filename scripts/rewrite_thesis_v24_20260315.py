# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (c) 2026 林志荣

from docx import Document
from pathlib import Path


INPUT = Path(r"\\wsl.localhost\Ubuntu\home\lzr\eiscore\docs\tmp_master_stable_for_v24.docx")
OUTPUT = Path(r"\\wsl.localhost\Ubuntu\home\lzr\eiscore\docs\毕业论文初稿v2.4_终修初稿_2026-03-15.docx")
OUTPUT_EN = Path(r"\\wsl.localhost\Ubuntu\home\lzr\eiscore\docs\thesis_draft_v2_4_final_draft_2026-03-15.docx")


REWRITE = {
    202: "Digitalization becomes difficult for Nanpai Food not because the company has never used software, but because one business record is still broken into several manual traces. Orders, stock movements, inspection files and workflow states do not stay on the same path for long. Some records remain on paper, some move through chat messages, and some are copied again into separate ledgers. Once one step is delayed, traceability becomes slower and coordination becomes unstable.",
    203: "This thesis therefore treats EISCore as a system core for resource-constrained manufacturers rather than a full enterprise suite. The practical target is narrower and more concrete: keep personnel governance, material circulation, application configuration and workflow linkage on a maintainable path. PostgreSQL carries the core data model together with part of the constraints and access boundaries, while PostgREST shortens the ordinary API chain. Vue 3 and qiankun keep a unified entry on the frontend and separate the modules that are already stable, including human resources, material management, the application center and mobile services.",
    204: "Within the present scope, the project has already stabilized two main business lines around personnel and materials. It also brings in BPMN-based workflow support, FlashBuilder, a controlled Agent Runtime, lightweight semantic enhancement and the cold-storage PDA mode. These parts are not presented as isolated highlights. They are connected through the same runtime, the same object boundaries and the same database-centered structure.",
    205: "Keywords: manufacturing SMEs; enterprise information core; micro-frontends; database-centered structure; lightweight semantic enhancement",
    272: "本文这里所说的“本体”，并不是要在系统外再搭一套独立知识图谱。它先解决表、字段和关系在不同模块里解释不一致的问题：流程配置这样用，权限判断又那样用，到了动态表单里口径可能再变一次。于是项目先给现有对象补上语义说明。现在这层语义已经进入 Agent 运行时，但运行时并不允许模型自己拼接口。它会先识别当前对象更接近数据表、流程状态还是语义关系，再转到白名单工具和对应接口。",
    377: "在接口与运行层，系统主要由 PostgREST、工作流运行组件和辅助运行组件构成。PostgREST 负责把数据库中的表、视图和函数映射成接口；工作流运行组件承载流程实例和任务流转；辅助运行组件则接住闪念应用相关的 Agent 接口、草稿同步、附件上传和工具编排。运行时目前已按白名单注册数据表、表格、流程、库存和语义相关工具，目的是把智能体调用限制在既有权限边界内。这样一来，数据表格应用、流程状态和轻量化语义能够被联动起来，但不会变成任意接口自由调用。",
    642: "应用中心模块当前已完成首页、数据应用配置、应用运行页面以及闪念应用构建器。用户可以在首页查看已有应用，在数据应用配置页维护字段与配置，在运行页查看应用效果，也可以通过 FlashBuilder 进入对话式草稿构建流程。当前这条链路已经具备草稿读取与保存、附件上传、预览、发布路由写入以及与 Agent Runtime 协同处理的基础能力。运行时提供的是受控工具，而不是开放式接口调用，因此闪念应用能够继续联动数据表格应用与语义增强结构，但仍然保持在既有权限边界内。"
}


def main():
    doc = Document(str(INPUT))
    changed = []
    for idx, text in REWRITE.items():
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].text = text
            changed.append(idx)

    doc.save(str(OUTPUT))
    doc.save(str(OUTPUT_EN))
    print(f"saved: {OUTPUT}")
    print(f"saved: {OUTPUT_EN}")
    print(f"changed: {len(changed)}")
    print(",".join(str(i) for i in changed))


if __name__ == "__main__":
    main()
