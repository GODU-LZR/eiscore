from __future__ import annotations

import re
from collections import defaultdict
from pathlib import Path

from docx import Document


DOCX_PATH = Path(r"C:\Users\Twist\Desktop\论文\主稿\毕业论文初稿v2.8_降风险终修版_2026-04-06.docx")
OUT_DIR = Path(r"C:\Users\Twist\Desktop\论文\清单说明\v28_降风险终修版_图表引用核查_2026-04-06")

FIG_RE = re.compile(r"图(\d+-\d+)")
TAB_RE = re.compile(r"表(\d+-\d+)")


def paragraph_texts(doc: Document) -> list[tuple[str, str]]:
    items: list[tuple[str, str]] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            items.append(("para", text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip()).strip()
                if text:
                    items.append(("cell", text))
    return items


def is_caption(text: str, prefix: str, number: str) -> bool:
    s = text.strip().replace(" ", "")
    return s.startswith(f"{prefix}{number}")


def main() -> None:
    doc = Document(str(DOCX_PATH))
    texts = paragraph_texts(doc)

    figures: dict[str, list[str]] = defaultdict(list)
    tables: dict[str, list[str]] = defaultdict(list)

    for kind, text in texts:
        for num in FIG_RE.findall(text):
            figures[num].append(text)
        for num in TAB_RE.findall(text):
            tables[num].append(text)

    lines: list[str] = ["# 图表正文引用核查", ""]

    lines.append("## 图片")
    fig_missing = []
    for num in sorted(figures.keys(), key=lambda x: tuple(map(int, x.split("-")))):
        refs = [t for t in figures[num] if not is_caption(t, "图", num)]
        caption_like = [t for t in figures[num] if is_caption(t, "图", num)]
        if refs:
            lines.append(f"- 图{num}: 已引用")
            lines.append(f"  - 标题/图题: {caption_like[0][:120] if caption_like else '未识别到独立图题'}")
            lines.append(f"  - 正文引用: {refs[0][:160]}")
        else:
            fig_missing.append(num)
            lines.append(f"- 图{num}: 未发现正文引用")
        lines.append("")

    lines.append("## 表格")
    tab_missing = []
    for num in sorted(tables.keys(), key=lambda x: tuple(map(int, x.split("-")))):
        refs = [t for t in tables[num] if not is_caption(t, "表", num)]
        caption_like = [t for t in tables[num] if is_caption(t, "表", num)]
        if refs:
            lines.append(f"- 表{num}: 已引用")
            lines.append(f"  - 标题/表题: {caption_like[0][:120] if caption_like else '未识别到独立表题'}")
            lines.append(f"  - 正文引用: {refs[0][:160]}")
        else:
            tab_missing.append(num)
            lines.append(f"- 表{num}: 未发现正文引用")
        lines.append("")

    lines.append("## 结论")
    lines.append(f"- 未发现正文引用的图片数量: {len(fig_missing)}")
    if fig_missing:
        lines.append(f"- 未引用图片: {', '.join('图' + n for n in fig_missing)}")
    lines.append(f"- 未发现正文引用的表格数量: {len(tab_missing)}")
    if tab_missing:
        lines.append(f"- 未引用表格: {', '.join('表' + n for n in tab_missing)}")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUT_DIR / "图表引用核查.md"
    out.write_text("\n".join(lines), encoding="utf-8")
    print(out)


if __name__ == "__main__":
    main()
