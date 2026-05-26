from __future__ import annotations

from pathlib import Path

from docx import Document


SRC = Path("/home/lzr/eiscore/docs/毕业论文最终主稿_稳妥版_2026-03-15.docx")
OUT = Path("/home/lzr/eiscore/docs/毕业论文最终主稿_终版_2026-03-15.docx")
OUT_EN = Path("/home/lzr/eiscore/docs/final_thesis_master_final_2026-03-15.docx")


PARA_REPLACEMENTS = {
    202: "In small and medium-sized manufacturing enterprises in Zhanjiang, the practical issue of digitalization is not simply whether software has been introduced, but whether order data, inventory records, personnel information, workflow states and quality documents can be connected and maintained in a consistent way. Guangdong Nanpai Food Co., Ltd., which is used as the business background of this thesis, shows several representative characteristics of this situation: raw materials circulate in multiple batches, some material attributes change from batch to batch, process adjustments are frequent, and many operational records still depend on manual transfer. Under these conditions, information fragmentation, weak traceability, unclear permission boundaries and inefficient collaboration are likely to appear. Large ERP systems are usually costly and heavy for this type of enterprise, while lightweight tools are easier to deploy but often weaker in customization and later expansion. This is the practical context in which a lighter and easier-to-maintain system core is needed.",
    203: "To deal with these problems, this thesis designs and implements EISCore as a lightweight enterprise information system core for small and medium-sized manufacturing enterprises. The system uses a database-centric structure: PostgreSQL carries the core data model together with part of the business constraints, and PostgREST is used to expose interfaces and reduce repetitive backend API development. On the frontend side, Vue 3 and qiankun are adopted to organize a unified entry with several coordinated modules, including human resource management, material management, the application center and mobile services. Within the current scope, the implemented system has already formed two relatively stable business lines around personnel governance and material circulation, covering employee archives, organizational structure, user-role management, attendance, material master data, warehouse management, inventory ledger, stock-in, stock-out and stock checking. BPMN-based workflow support and a lightweight semantic enhancement mechanism are also introduced so that workflow states, permission control and later extension needs can be connected more consistently.",
    204: "The thesis is organized into nine chapters. Chapter 1 discusses the research background, related studies and the research focus of this work. Chapter 2 introduces the main technologies and the runtime environment used in EISCore. Chapter 3 derives system requirements from the enterprise scenario, business roles and operational problems. Chapter 4 presents the overall system design, and Chapter 5 explains the database structure together with the semantic enhancement design. Chapter 6 describes the detailed design of key parts such as permission control, inventory processing, workflow linkage and semantic support. Chapter 7 shows the implementation results, Chapter 8 reports the testing work, and Chapter 9 concludes the thesis and discusses the current limitations together with possible future improvements.",
    205: "Keywords: small and medium-sized manufacturing enterprises; enterprise informatization system; micro-frontends; database-centric architecture; lightweight semantic enhancement",
}


def main() -> None:
    doc = Document(str(SRC))
    changed = 0
    for idx, new_text in PARA_REPLACEMENTS.items():
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].text = new_text
            changed += 1
    doc.save(str(OUT))
    doc.save(str(OUT_EN))
    print(f"saved: {OUT}")
    print(f"saved: {OUT_EN}")
    print(f"changed paragraphs: {changed}")


if __name__ == "__main__":
    main()
