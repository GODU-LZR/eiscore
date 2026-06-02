# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (c) 2026 林志荣

from __future__ import annotations

from pathlib import Path

from docx import Document


SRC = Path(r"\\wsl.localhost\Ubuntu\home\lzr\eiscore\docs\毕业论文最终主稿_终版_补闪念语义冷库模式_2026-03-15.docx")
OUT = Path(r"\\wsl.localhost\Ubuntu\home\lzr\eiscore\docs\毕业论文最终主稿_终版_补闪念语义冷库模式_润色版_2026-03-15.docx")
OUT_EN = Path(r"\\wsl.localhost\Ubuntu\home\lzr\eiscore\docs\final_thesis_master_final_with_flash_semantic_coldmode_tuned_2026-03-15.docx")


PARA_REPLACEMENTS = {
    233: "围绕前述问题，本文对 EISCore 的研究与实现主要集中在六个方面。其一，围绕人事管理和物料管理两条主线建立核心业务结构，把员工、组织、角色、物料主数据、库存台账、批次和出入库链路组织起来。其二，在后端采用数据库中心架构，由 PostgreSQL 承载核心数据和部分业务逻辑，并借助 PostgREST 压缩常规接口层。其三，在前端以微前端方式组织基座、人事、物料、应用中心和移动端，使模块边界保持清晰。其四，在应用中心中落地了面向草稿生成与应用发布的闪念应用构建器，把提示输入、附件、草稿、预览和发布路由接成一条辅助开发链路。其五，在这条链路上继续引入受控的智能体能力注册机制，把数据表结构保障、表格数据读写、流程状态访问、本体关系查询和语义增强分别组织为带有 domain、object、intent 的工具，再交给运行时在白名单和权限边界内调用。其六，在平台层补入 BPMN 流程建模与轻量化语义机制，分别支撑审批状态联动和表、字段、关系结构的可解释表达。",
    252: "除基础组件外，系统还结合页面需求引入了 ECharts、html2canvas、Monaco Editor、Markdown-it 和 Mammoth 等工具。它们分别服务于统计图展示、页面导出、配置编辑、闪念应用草稿构建以及文本与文档处理场景。结合当前实现，这些工具并不是散落在不同页面里的附属插件。Monaco Editor 主要用于 FlashBuilder 中的草稿查看和编辑，Mammoth 与 Markdown-it 负责把需求文档和文本输入转成运行时可消费的内容，AG Grid 则继续承担数据表格应用中的列表承载任务。这样一来，闪念应用构建器、数据应用配置页和应用运行页共用了一套较稳定的交互支撑链路。",
    272: "本文所说的“本体”更接近一种面向业务表结构的语义增强方式，而不是额外建设独立知识图谱。它的任务是给现有表、字段和表间关系补充可解释的语义描述，使流程、权限和动态配置在同一套数据结构上保持更一致的表达口径。就当前项目而言，这套语义描述已经开始进入 Agent 运行时的能力建模过程：运行时并不允许模型自由拼接接口，而是先把数据表、流程、库存和本体关系组织成带有 domain、object、intent 的 semantic tool，再据此映射到允许调用的接口。",
}


def main() -> None:
    doc = Document(str(SRC))
    changed = 0
    for idx, text in PARA_REPLACEMENTS.items():
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].text = text
            changed += 1
    doc.save(str(OUT))
    doc.save(str(OUT_EN))
    print(f"saved: {OUT}")
    print(f"saved: {OUT_EN}")
    print(f"changed paragraphs: {changed}")


if __name__ == "__main__":
    main()
