# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (c) 2026 林志荣

from docx import Document
from pathlib import Path


INPUT = Path(r"\\wsl.localhost\Ubuntu\home\lzr\eiscore\docs\tmp_thesis_v21_current.docx")
OUTPUT = Path(r"\\wsl.localhost\Ubuntu\home\lzr\eiscore\docs\毕业论文初稿v2.2_手册精修版_2026-03-15.docx")
OUTPUT_EN = Path(r"\\wsl.localhost\Ubuntu\home\lzr\eiscore\docs\thesis_draft_v2_2_handbook_refined_2026-03-15.docx")


REWRITE = {
    197: "在南派食品的调研里，最麻烦的并不是少一个页面，而是一批业务记录被拆散后很难再接回去。销售单、领料单、检验附件和库存变动分别留在纸单、群消息和分散台账里，任何一处补录不及时，后面追原料批次、仓位、补料和出库去向都会变慢。对湛江不少中小制造企业来说，问题大多也是这样来的：预算有限，运维人手不多，现场规则却一直在变。系统做得太重，后面难养；做得太散，数据链路又接不起来。",
    198: "本文据此设计并实现 EISCore。它没有先铺一层更厚的应用服务，而是先把核心数据、部分业务约束和访问边界收回 PostgreSQL 与 PostgREST 这一条链路，再由微前端基座承载人事、物料、应用中心和移动端。数据库、接口、运行时与前端进程统一交给 Docker Compose 和 PM2 组织。这样做并不追求花哨。先把人事治理、库存台账、应用配置、流程联动和语义增强这几条容易失控的链路收稳，后续扩展才有基础。",
    202: "At Nanpai Food, the difficult part is not the lack of software names on paper. Trouble starts when one business record is split into several manual traces. Orders, lot notes, inspection files and stock movements still move across paper forms, chat messages and separate ledgers. Once one step lags behind, traceability slows down and coordination becomes fragile. This is the kind of constraint many resource-constrained manufacturers in Zhanjiang still face.",
    203: "EISCore is therefore built as a lightweight system core rather than a full enterprise suite. The project does not begin with a heavy service layer. Core data, part of the business rules and access boundaries stay close to PostgreSQL, while PostgREST removes repetitive API work. On the client side, Vue 3 and qiankun keep a unified entry and separate the modules that have already become stable: human resources, material management, the application center and mobile services. Workflow support, FlashBuilder, controlled agent tools, semantic enhancement and the cold-storage PDA mode are then attached to the same runtime chain.",
    204: "The thesis follows the actual engineering path of the project. Chapter 1 explains the problem background and related studies. Chapter 2 introduces the technical stack and runtime environment. Chapter 3 derives requirements from the Nanpai scenario. Chapters 4 to 6 move from overall design to database design and detailed design. Chapter 7 presents the implemented functions, Chapter 8 reports the tests, and Chapter 9 summarizes the work together with the remaining limitations and future extensions.",
    223: "现有国内研究已经给出不少可用方法，但把模块建设、流程联动、权限治理和轻量化语义同时放进一个可运行平台里的案例仍然不多。问题不完全在理论。真正难的是整合：这些能力放到资源有限的企业现场后，谁先落地、谁负责边界、谁和谁共用一套对象口径。围绕这一层，现有样本还比较少，本文就在这里继续展开。",
    230: "综合来看，现有研究覆盖了数字化转型、流程建模和语义表达等多个方向，可一旦落到中小制造企业核心业务系统，短板往往出现在整合环节。很多成果按主题分别推进，业务模块、流程协同、权限治理和语义表达没有一起设计；不少成熟方案又更适用于大型企业或高标准化场景。换到“小批量、多批次、频繁变更”的地方制造企业里，问题会变得很具体：字段要改、状态规则要调、表单要跟着变，这时厚重服务层带来的维护成本就很快显出来了。",
    272: "本文这里说的“本体”，并不是要在系统旁边再建一套独立知识图谱。它更像给现有业务表结构补一层能解释、也能复用的语义。表、字段和表间关系先被说明白，流程、权限和动态配置才能尽量沿用同一套对象口径。现在这层语义已经进入 Agent 运行时，不过运行时并不会让模型自由拼接口。它会先判断当前对象更接近数据表、流程状态还是语义关系，然后再落到白名单工具，例如 `flash.data.table.ensure`、`flash.workflow.instance.start` 或 `flash.ontology.semantic.enrich`。",
    278: "本章涉及的技术并不是并列摆放的名词。它们围绕同一条运行链路分工：前端承载模块交互，数据库保存核心数据并收住约束，PostgREST 缩短接口链路，BPMN 负责过程型业务，轻量化语义补对象解释，Docker Compose 与 PM2 则把这些部分组织进同一套可复现环境。对本文来说，关键不是技术名词多不多，而是这套结构能不能在资源有限的企业环境里继续维护。",
    295: "这一段流程里最难的并不是下单本身，而是状态始终跟不齐。销售合同到销售订单之间缺少高效传递机制，客户、物料和数量等信息往往要重复录入；订单状态在销售、计划和仓库之间的传递也不够透明，销售人员很难及时知道“是否排产、是否完工、是否可发货”。系统因此需要支持从销售合同到销售订单的一键下推，并建立订单状态可视化机制。只有这样，销售、PMC 和仓库才可能围绕同一张订单形成一致判断。",
    359: "本章基于广东南派食品有限公司的调研资料，收拢了企业现状、角色边界、业务流程、功能需求和非功能需求。补料、质检延迟、客户物料映射、批次与保质期管理、纸质附件留存、冷库离线盘点以及细粒度权限控制等问题都被保留下来。它们不是零散细节，而是后续设计必须回应的真实约束。后面的总体设计和数据库设计，都围绕这些约束继续展开。",
    376: "在前端表现层，系统仍由基座应用收一个统一入口。人事子应用把员工、组织、岗位和考勤放在同一侧维护；物料子应用围绕物料主数据、仓储台账、出入库和批次管理展开；应用中心继续承接数据应用配置、闪念应用构建器、流程配置、审批中心和本体关系工作台。移动端不追求复制完整桌面能力，它更偏向现场：盘点、扫码、标签打印，以及冷库弱网场景下的补充入口。",
    434: "南派食品仓储并不是单一平面结构，实际需要同时区分仓库、库区和库位。`scm.warehouses` 因此通过 `parent_id` 组织成树形层级，让仓位定位、库存展示和后续盘点都还能沿着同一套空间结构走。",
    436: "库存批次表 `scm.inventory_batches` 是食品行业批次管理的关键表。它不是在库存表旁边随手补一个批次字段，而是把物料、仓库、批次号、可用数量、锁定数量、生产日期、过期日期和状态单独收成一个对象。这样做的原因很现实：对南派食品而言，批次是追溯链能不能接起来的支点。",
    438: "库存流水表 `scm.inventory_transactions` 记录所有库存变动，包括采购入库、生产领料、补料出库、生产入库、销售出库和盘点调整。它通过 `material_id`、`batch_id` 和 `warehouse_id` 接回物料、批次与仓库，同时保留交易类型、数量、关联单据和操作时间等过程信息。",
    439: "这张表最直接的作用有两个。第一，批次数量变化不再只有结果，还能回查来源；第二，静态库存状态和动态业务动作被接到了同一条链路上，后续追溯和业务分析才有过程数据可看。没有库存流水，库存台账就只剩下一个静态结果。",
    605: "轻量化语义模块并不是单独摆在应用中心里看图用的。当前它已经和三个方向接上：数据应用配置会引用语义信息辅助字段展示；流程配置和状态映射借它保持对象解释一致；敏感字段和重要关系也可以与权限控制兼容表达。语义因此不再停留在说明层，而是开始进入支撑层。",
    611: "这一节完成后，轻量化语义不再只是论文里的概念补充。它已经为应用配置、流程联动和对象解释提供了同一层底层表达，这也是本文区别于一般信息系统实现的一处关键落点。",
    613: "本章围绕用户与权限控制、库存入库与出库、流程审批与状态映射以及轻量化语义模块展开详细设计。重点不在罗列页面，而在把前端页面、数据库函数、流程控制与语义治理接成一条能够运行的实现链路。下一章的实现与测试，都会回到这条链路上验证。",
    623: "在人事管理模块中，员工档案、组织结构和岗位管理已经形成可用页面。员工列表负责维护基础信息，组织结构页面负责部门层级，用户管理页面把账号、岗位和角色关系收在一起。这样一来，人员信息不再散落在多张 Excel 或临时台账里；后续做权限判断和流程候选人分配时，也有统一入口可回查。",
    624: "这意味着部门、岗位和用户关系不再只是静态主数据，而开始承担页面访问控制和流程分派的基础角色。",
    684: "23 项检查全部通过。这个结果不能说明系统已经完全成熟，但至少把论文里最关键的几条链路坐实了：登录鉴权、子应用访问、核心数据接口、流程定义读取、辅助运行组件和主要进程都能稳定工作。围绕 Agent Runtime 的表格工具、本体关系查询工具和语义增强工具也完成了基础验证，因此“智能体—数据表格应用—轻量化语义”这条链路已经不只停留在设计层。移动端侧的冷库模式缓存、本地待处理记录和离线盘点入口同样通过了阶段性验证。",
    725: "本文讨论的不是抽象口号，而是南派食品这类企业每天都会碰到的摩擦：信息分散、库存台账不清、批次追溯慢、权限边界不清，后续改动一多，系统维护成本还会继续往上走。围绕这些问题，论文完成了需求分析、总体设计、数据库设计、详细设计、系统实现和阶段性测试，并把结果落实为可运行的原型。",
    726: "从当前完成情况看，系统骨架已经立住。前端由基座、人事、物料、应用中心和移动端协同运行；后端以数据库中心架构承载主要数据结构、部分业务约束和权限边界，流程配置与审批联动也完成了阶段性落地。数据库、接口与运行时服务通过 Docker Compose 统一组织，多个前端模块通过 PM2 管理运行，这让部署和后续维护方式至少是清楚的。",
    729: "应用中心和流程能力的落地也不是零散页面拼接。数据应用配置、闪念应用构建器、流程设计、审批中心和应用运行页面已经打通到同一条运行链路上；流程定义、流程实例、任务分派和状态映射则为这条链路提供底层结构。更关键的是，FlashBuilder 负责草稿与发布，Agent Runtime 通过白名单工具接入数据表、流程和本体关系，数据表格应用与语义结构再把对象和字段基础补齐。",
    730: "另一项更有辨识度的工作，是把轻量化语义真正放进系统实现过程，而不是停留在术语说明层。表级语义、列级语义、中文化回填和本体关系工作台已经落地，语义对象也开始进入 Agent 工具组织。它不追求复杂推理，先解决的是一个更实际的问题：让流程、权限、动态配置和辅助构建尽量共用一套对象解释。",
    731: "在测试与验证方面，本文基于当前项目运行环境，对登录鉴权、微前端页面访问、核心数据接口、流程相关数据读取、闪念应用相关运行链路、辅助运行组件和运行状态等内容进行了阶段性测试。结果表明，系统当前主干功能运行稳定，关键页面和关键链路能够满足论文阶段的展示与验证需要。",
    732: "整体来看，本文已经完成了从需求分析到实现验证的主要链条，也证明了数据库中心架构、微前端拆分、流程支撑和轻量化语义机制在中小制造企业场景下具备现实可行性。以南派食品为例，这套系统虽然尚未覆盖全部业务方向，但已经把最核心的几条链路收稳，并为后续扩展留下了清晰基础。",
}


def main():
    doc = Document(str(INPUT))
    changed = []
    for idx, text in REWRITE.items():
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].text = text
            changed.append(idx)

    doc.save(str(OUTPUT))
    doc.save(str(OUTPUT_EN))
    print(f"saved: {OUTPUT}")
    print(f"saved: {OUTPUT_EN}")
    print(f"changed: {len(changed)}")
    print(",".join(str(i) for i in changed))


if __name__ == "__main__":
    main()
